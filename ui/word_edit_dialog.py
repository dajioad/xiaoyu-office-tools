# word_edit_dialog.py
import os
import json
import base64
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QTextEdit,
    QFileDialog, QMessageBox, QLabel, QApplication, QWidget,
    QToolBar, QStatusBar, QMenuBar, QMenu,
    QCheckBox, QLineEdit, QColorDialog, QInputDialog, QProgressDialog
)
from PySide6.QtCore import Qt, QTimer, QSize
from PySide6.QtGui import QTextDocument
from PySide6.QtGui import (
    QFont, QAction, QKeySequence, QTextCursor, QColor, QIcon
)

import qtawesome as qta

try:
    from main_window import THEME_FONT
except ImportError:
    THEME_FONT = "Segoe UI"

ICON_SIZE_MEDIUM = QSize(20, 20)


class ColorPickerDialog(QDialog):
    def __init__(self, parent=None, initial_color=None):
        super().__init__(parent)
        self.selected_color = initial_color or QColor("#000000")
        self.setWindowTitle("选择颜色")
        self.setMinimumSize(360, 400)
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(16, 16, 16, 16)

        preset_label = QLabel("常用颜色:")
        preset_label.setStyleSheet("font-weight: bold; font-size: 13px;")
        layout.addWidget(preset_label)

        preset_colors = [
            ("#000000", "黑色"), ("#FFFFFF", "白色"), ("#FF0000", "红色"), ("#00FF00", "绿色"),
            ("#0000FF", "蓝色"), ("#FFFF00", "黄色"), ("#FF00FF", "品红"), ("#00FFFF", "青色"),
            ("#FFA500", "橙色"), ("#800080", "紫色"), ("#00FF7F", "春绿"), ("#DC143C", "深红"),
            ("#4169E1", "皇家蓝"), ("#FF69B4", "粉红"), ("#A52A2A", "棕色"), ("#808080", "灰色"),
            ("#C0C0C0", "银色"), ("#FFD700", "金色"), ("#008080", "深青"), ("#800000", "栗色"),
            ("#FFC0CB", "浅粉"), ("#87CEEB", "天蓝"), ("#90EE90", "浅绿"), ("#DDA0DD", "梅红"),
        ]

        color_grid = QWidget()
        color_grid_layout = QVBoxLayout(color_grid)
        color_grid_layout.setSpacing(4)
        color_grid_layout.setContentsMargins(0, 0, 0, 0)

        for i in range(0, len(preset_colors), 6):
            row_widget = QWidget()
            row_layout = QHBoxLayout(row_widget)
            row_layout.setSpacing(4)
            row_layout.setContentsMargins(0, 0, 0, 0)
            for j in range(6):
                if i + j < len(preset_colors):
                    color_hex, color_name = preset_colors[i + j]
                    color_btn = QPushButton()
                    color_btn.setFixedSize(40, 32)
                    color_btn.setStyleSheet(f"""
                        QPushButton {{
                            background-color: {color_hex};
                            border: 1px solid #d0d0d0;
                            border-radius: 4px;
                        }}
                        QPushButton:hover {{
                            border: 2px solid #409eff;
                        }}
                    """)
                    color_btn.setToolTip(f"{color_name} ({color_hex})")
                    color_btn.clicked.connect(lambda checked, c=color_hex: self._on_color_selected(c))
                    row_layout.addWidget(color_btn)
                else:
                    row_layout.addWidget(QWidget())
            color_grid_layout.addWidget(row_widget)
        layout.addWidget(color_grid)

        layout.addWidget(QLabel("或选择自定义颜色:"))
        custom_row = QHBoxLayout()
        self.custom_color_btn = QPushButton("  选择颜色  ")
        self.custom_color_btn.setIcon(qta.icon('fa5s.palette', color='white'))
        self.custom_color_btn.clicked.connect(self._show_color_dialog)
        self.color_preview = QWidget()
        self.color_preview.setFixedSize(60, 32)
        self.color_preview.setStyleSheet(f"""
            QWidget {{
                background-color: {self.selected_color.name()};
                border: 1px solid #d0d0d0;
                border-radius: 4px;
            }}
        """)
        custom_row.addWidget(self.custom_color_btn)
        custom_row.addWidget(self.color_preview)
        custom_row.addStretch()
        layout.addLayout(custom_row)

        btn_row = QHBoxLayout()
        ok_btn = QPushButton("确定")
        ok_btn.setStyleSheet("""
            QPushButton {
                padding: 8px 24px;
                background-color: #409eff;
                color: white;
                border: none;
                border-radius: 6px;
            }
            QPushButton:hover { background-color: #66b1ff; }
        """)
        ok_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton("取消")
        cancel_btn.setStyleSheet("""
            QPushButton {
                padding: 8px 24px;
                background-color: #f5f5f5;
                color: #606266;
                border: 1px solid #dcdfe6;
                border-radius: 6px;
            }
            QPushButton:hover { background-color: #e4e4e4; }
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addStretch()
        btn_row.addWidget(ok_btn)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

    def _on_color_selected(self, color_hex):
        self.selected_color = QColor(color_hex)
        self.color_preview.setStyleSheet(f"""
            QWidget {{
                background-color: {color_hex};
                border: 1px solid #d0d0d0;
                border-radius: 4px;
            }}
        """)

    def _show_color_dialog(self):
        color = QColorDialog.getColor(self.selected_color, self, "选择颜色")
        if color.isValid():
            self.selected_color = color
            self.color_preview.setStyleSheet(f"""
                QWidget {{
                    background-color: {color.name()};
                    border: 1px solid #d0d0d0;
                    border-radius: 4px;
                }}
            """)

    def get_color(self):
        return self.selected_color


class FindReplaceDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setWindowTitle("查找与替换")
        self.setMinimumWidth(400)
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(16, 16, 16, 16)
        
        self._apply_theme()
        
        find_layout = QHBoxLayout()
        find_layout.addWidget(QLabel("查找内容:"))
        self.find_edit = QLineEdit()
        self.find_edit.setStyleSheet(f"""
            QLineEdit {{
                border: 1px solid {self._border_color};
                border-radius: 6px;
                padding: 6px 10px;
                background-color: {self._bg_color};
                color: {self._text_color};
            }}
        """)
        find_layout.addWidget(self.find_edit)
        layout.addLayout(find_layout)
        
        replace_layout = QHBoxLayout()
        replace_layout.addWidget(QLabel("替换为:"))
        self.replace_edit = QLineEdit()
        self.replace_edit.setStyleSheet(f"""
            QLineEdit {{
                border: 1px solid {self._border_color};
                border-radius: 6px;
                padding: 6px 10px;
                background-color: {self._bg_color};
                color: {self._text_color};
            }}
        """)
        replace_layout.addWidget(self.replace_edit)
        layout.addLayout(replace_layout)
        
        options_layout = QHBoxLayout()
        options_layout.setAlignment(Qt.AlignLeft)
        
        self.case_sensitive = QCheckBox("区分大小写")
        self.whole_word = QCheckBox("全字匹配")
        self.whole_word.setToolTip("注：此选项仅对英文单词有效，中文请使用精确匹配")
        
        checkbox_style = f"""
            QCheckBox {{
                color: {self._text_color};
                font-size: 13px;
            }}
        """
        self.case_sensitive.setStyleSheet(checkbox_style)
        self.whole_word.setStyleSheet(checkbox_style)
        
        # 添加提示标签
        self.whole_word_hint = QLabel("<span style='color: #909399; font-size: 11px;'>（仅英文）</span>")
        
        options_layout.addWidget(self.case_sensitive)
        options_layout.addWidget(self.whole_word)
        options_layout.addWidget(self.whole_word_hint)
        options_layout.addStretch()
        layout.addLayout(options_layout)
        
        btn_layout = QHBoxLayout()
        self.find_next_btn = QPushButton("查找下一个")
        self.replace_btn = QPushButton("替换")
        self.replace_all_btn = QPushButton("全部替换")
        
        btn_style = f"""
            QPushButton {{
                padding: 6px 16px;
                background-color: {self._primary_color};
                color: white;
                border: none;
                border-radius: 6px;
                font-size: 13px;
            }}
            QPushButton:hover {{
                background-color: {self._primary_hover};
            }}
            QPushButton:pressed {{
                background-color: {self._primary_color};
            }}
        """
        self.find_next_btn.setStyleSheet(btn_style)
        self.replace_btn.setStyleSheet(btn_style)
        self.replace_all_btn.setStyleSheet(btn_style)
        
        btn_layout.addWidget(self.find_next_btn)
        btn_layout.addWidget(self.replace_btn)
        btn_layout.addWidget(self.replace_all_btn)
        layout.addLayout(btn_layout)
        
        self.find_next_btn.clicked.connect(self.find_next)
        self.replace_btn.clicked.connect(self.replace_current)
        self.replace_all_btn.clicked.connect(self.replace_all)
    
    def _apply_theme(self):
        """应用主题样式"""
        theme = {
            'bg': '#ffffff',
            'card': '#ffffff',
            'text': '#303133',
            'border': '#dcdfe6',
            'primary': '#409eff',
            'primary_hover': '#66b1ff'
        }
        
        if self.parent and hasattr(self.parent, 'theme_manager'):
            try:
                theme = self.parent.theme_manager.get_theme()
            except:
                pass
        
        self._bg_color = theme.get('bg', '#ffffff')
        self._text_color = theme.get('text', '#303133')
        self._border_color = theme.get('border', '#dcdfe6')
        self._primary_color = theme.get('primary', '#409eff')
        self._primary_hover = theme.get('primary_hover', '#66b1ff')
        
        self.setStyleSheet(f"background-color: {theme.get('card', '#ffffff')};")
    
    def find_next(self):
        text = self.find_edit.text()
        if not text:
            QMessageBox.warning(self, "提示", "请输入查找内容")
            return False
        doc = self.parent.text_edit.document()
        cursor = self.parent.text_edit.textCursor()
        flags = QTextDocument.FindFlag(0)
        if self.case_sensitive.isChecked():
            flags |= QTextDocument.FindCaseSensitively
        if self.whole_word.isChecked():
            flags |= QTextDocument.FindWholeWords
        new_cursor = doc.find(text, cursor, flags)
        if not new_cursor.isNull():
            self.parent.text_edit.setTextCursor(new_cursor)
            return True
        else:
            cursor.movePosition(QTextCursor.Start)
            new_cursor = doc.find(text, cursor, flags)
            if not new_cursor.isNull():
                self.parent.text_edit.setTextCursor(new_cursor)
                QMessageBox.information(self, "提示", "已循环到开头")
                return True
            else:
                QMessageBox.information(self, "提示", "找不到内容")
                return False
    
    def replace_current(self):
        if not self.find_next():
            return
        cursor = self.parent.text_edit.textCursor()
        if cursor.hasSelection():
            cursor.insertText(self.replace_edit.text())
            self.parent.modified = True
    
    def replace_all(self):
        search_text = self.find_edit.text()
        replace_text = self.replace_edit.text()
        if not search_text:
            return
        doc = self.parent.text_edit.document()
        cursor = QTextCursor(doc)
        cursor.movePosition(QTextCursor.Start)
        flags = QTextDocument.FindFlag(0)
        if self.case_sensitive.isChecked():
            flags |= QTextDocument.FindCaseSensitively
        if self.whole_word.isChecked():
            flags |= QTextDocument.FindWholeWords
        count = 0
        while True:
            cursor = doc.find(search_text, cursor, flags)
            if cursor.isNull():
                break
            cursor.insertText(replace_text)
            count += 1
        if count > 0:
            self.parent.modified = True
            QMessageBox.information(self, "完成", f"共替换 {count} 处")
        else:
            QMessageBox.information(self, "提示", "没有找到匹配内容")


class WordContextMenuHelper:
    def __init__(self, text_edit, parent_dialog):
        self.text_edit = text_edit
        self.parent = parent_dialog

    def create_context_menu(self, pos):
        menu = QMenu(self.parent)

        # 基本编辑
        cut_action = QAction("剪切", self.parent)
        cut_action.setIcon(qta.icon('fa5s.cut', color='#606266'))
        cut_action.setShortcut(QKeySequence.Cut)
        cut_action.triggered.connect(self.text_edit.cut)
        menu.addAction(cut_action)

        copy_action = QAction("复制", self.parent)
        copy_action.setIcon(qta.icon('fa5s.copy', color='#606266'))
        copy_action.setShortcut(QKeySequence.Copy)
        copy_action.triggered.connect(self.text_edit.copy)
        menu.addAction(copy_action)

        paste_action = QAction("粘贴", self.parent)
        paste_action.setIcon(qta.icon('fa5s.paste', color='#606266'))
        paste_action.setShortcut(QKeySequence.Paste)
        paste_action.triggered.connect(self.text_edit.paste)
        menu.addAction(paste_action)

        delete_action = QAction("删除", self.parent)
        delete_action.setIcon(qta.icon('fa5s.trash-alt', color='#606266'))
        delete_action.setShortcut(QKeySequence.Delete)
        delete_action.triggered.connect(lambda: self.text_edit.textCursor().removeSelectedText())
        menu.addAction(delete_action)
        menu.addSeparator()

        select_all_action = QAction("全选", self.parent)
        select_all_action.setIcon(qta.icon('fa5s.check-double', color='#606266'))
        select_all_action.setShortcut(QKeySequence.SelectAll)
        select_all_action.triggered.connect(self.text_edit.selectAll)
        menu.addAction(select_all_action)
        menu.addSeparator()

        # 文本格式化
        format_menu = QMenu("文本格式化", menu)
        bold_action = QAction("加粗", format_menu)
        bold_action.setIcon(qta.icon('fa5s.bold', color='#606266'))
        bold_action.setShortcut(QKeySequence("Ctrl+B"))
        bold_action.triggered.connect(self.toggle_bold)
        format_menu.addAction(bold_action)

        italic_action = QAction("斜体", format_menu)
        italic_action.setIcon(qta.icon('fa5s.italic', color='#606266'))
        italic_action.setShortcut(QKeySequence("Ctrl+I"))
        italic_action.triggered.connect(self.toggle_italic)
        format_menu.addAction(italic_action)

        underline_action = QAction("下划线", format_menu)
        underline_action.setIcon(qta.icon('fa5s.underline', color='#606266'))
        underline_action.setShortcut(QKeySequence("Ctrl+U"))
        underline_action.triggered.connect(self.toggle_underline)
        format_menu.addAction(underline_action)
        format_menu.addSeparator()
        color_action = QAction("文字颜色...", format_menu)
        color_action.setIcon(qta.icon('fa5s.palette', color='#606266'))
        color_action.triggered.connect(self.change_text_color)
        format_menu.addAction(color_action)
        menu.addMenu(format_menu)

        # 符号处理（核心）
        symbol_menu = self.parent.create_symbol_menu()
        if symbol_menu:
            menu.addMenu(symbol_menu)
        menu.addSeparator()

        # 插入菜单（仅图片）
        insert_menu = QMenu("插入", menu)
        image_action = QAction("插入图片...", insert_menu)
        image_action.setIcon(qta.icon('fa5s.image', color='#606266'))
        image_action.setToolTip("插入图片 (Ctrl+G)")
        image_action.triggered.connect(self.parent.insert_image)
        insert_menu.addAction(image_action)
        menu.addMenu(insert_menu)

        # 文本对齐
        align_menu = QMenu("文本对齐", menu)
        left_align = QAction("左对齐", align_menu)
        left_align.setIcon(qta.icon('fa5s.align-left', color='#606266'))
        left_align.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignLeft))
        align_menu.addAction(left_align)
        center_align = QAction("居中对齐", align_menu)
        center_align.setIcon(qta.icon('fa5s.align-center', color='#606266'))
        center_align.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignCenter))
        align_menu.addAction(center_align)
        right_align = QAction("右对齐", align_menu)
        right_align.setIcon(qta.icon('fa5s.align-right', color='#606266'))
        right_align.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignRight))
        align_menu.addAction(right_align)
        justify_align = QAction("两端对齐", align_menu)
        justify_align.setIcon(qta.icon('fa5s.align-justify', color='#606266'))
        justify_align.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignJustify))
        align_menu.addAction(justify_align)
        menu.addMenu(align_menu)

        return menu

    def toggle_bold(self):
        fmt = self.text_edit.currentCharFormat()
        fmt.setFontWeight(QFont.Bold if fmt.fontWeight() != QFont.Bold else QFont.Normal)
        self.text_edit.setCurrentCharFormat(fmt)
        self.parent.modified = True

    def toggle_italic(self):
        fmt = self.text_edit.currentCharFormat()
        fmt.setFontItalic(not fmt.fontItalic())
        self.text_edit.setCurrentCharFormat(fmt)
        self.parent.modified = True

    def toggle_underline(self):
        fmt = self.text_edit.currentCharFormat()
        fmt.setFontUnderline(not fmt.fontUnderline())
        self.text_edit.setCurrentCharFormat(fmt)
        self.parent.modified = True

    def change_text_color(self):
        current_fmt = self.text_edit.currentCharFormat()
        initial_color = current_fmt.foreground().color() if current_fmt.foreground().color().isValid() else QColor("#000000")
        color_dlg = ColorPickerDialog(self.parent, initial_color)
        if color_dlg.exec():
            color = color_dlg.get_color()
            fmt = self.text_edit.currentCharFormat()
            fmt.setForeground(color)
            self.text_edit.setCurrentCharFormat(fmt)
            self.parent.modified = True


class WordEditDialog(QDialog):
    def __init__(self, parent=None, is_pro=True, theme_manager=None, config_manager=None, main_window=None):
        super().__init__(parent)
        self.is_pro = is_pro
        self.theme_manager = theme_manager
        self.config_manager = config_manager
        self.main_window = main_window
        self.setWindowTitle("文本处理 - 文档编辑")
        self.setMinimumSize(1000, 700)
        self.resize(1200, 800)
        
        # 设置窗口图标
        import os
        from pathlib import Path
        app_dir = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        icon_path = app_dir / "小雨办公工具.png"
        if icon_path.exists():
            self.setWindowIcon(QIcon(str(icon_path)))

        self.current_path = None
        self.modified = False
        self.recent_files = []
        # split_mode 支持
        if self.main_window and hasattr(self.main_window, 'split_mode'):
            self.split_mode = self.main_window.split_mode
        else:
            self.split_mode = 0
        self._load_recent_files()          # 持久化加载
        self._is_first_show = True

        self.init_ui()
        self.setup_shortcuts()

        if not self.is_pro:
            self._set_readonly_mode()

        if self.theme_manager:
            self.apply_theme()

    def update_split_mode(self, split_mode):
        """由主窗口调用，更新分隔模式"""
        self.split_mode = split_mode

    # ---------- 持久化最近文件 ----------
    def _load_recent_files(self):
        if self.config_manager:
            recent_str = self.config_manager.get_config("Word", "recent_files", "[]")
            try:
                self.recent_files = json.loads(recent_str)
            except:
                self.recent_files = []
        else:
            self.recent_files = []

    def _save_recent_files(self):
        if self.config_manager:
            self.config_manager.set_config("Word", "recent_files", json.dumps(self.recent_files))

    def _add_recent_file(self, path):
        if path in self.recent_files:
            self.recent_files.remove(path)
        self.recent_files.insert(0, path)
        if len(self.recent_files) > 10:
            self.recent_files.pop()
        self._save_recent_files()

    # ---------- UI 初始化 ----------
    def init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # 极简菜单栏（文件菜单）
        menubar = QMenuBar()
        file_menu = menubar.addMenu("文件")
        new_act = QAction("新建", self)
        new_act.setShortcut(QKeySequence.New)
        new_act.triggered.connect(self.new_document)
        new_act.setEnabled(self.is_pro)
        file_menu.addAction(new_act)

        open_act = QAction("打开...", self)
        open_act.setShortcut(QKeySequence.Open)
        open_act.triggered.connect(self.open_file)
        file_menu.addAction(open_act)

        save_act = QAction("保存", self)
        save_act.setShortcut(QKeySequence.Save)
        save_act.triggered.connect(self.save_file)
        save_act.setEnabled(self.is_pro)
        file_menu.addAction(save_act)

        save_as_act = QAction("另存为...", self)
        save_as_act.setShortcut(QKeySequence.SaveAs)
        save_as_act.triggered.connect(self.save_file_as)
        save_as_act.setEnabled(self.is_pro)
        file_menu.addAction(save_as_act)

        file_menu.addSeparator()
        exit_act = QAction("退出", self)
        exit_act.triggered.connect(self.close)
        file_menu.addAction(exit_act)
        self.menubar = menubar   # 新增：保存菜单栏引用
        main_layout.addWidget(menubar)

        # 精简工具栏
        toolbar = QToolBar()
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(20, 20))

        open_btn = QAction(qta.icon('fa5s.folder-open', color='#409eff'), "打开", self)
        open_btn.triggered.connect(self.open_file)
        toolbar.addAction(open_btn)

        save_btn = QAction(qta.icon('fa5s.save', color='#409eff'), "保存", self)
        save_btn.triggered.connect(self.save_file)
        save_btn.setEnabled(self.is_pro)
        toolbar.addAction(save_btn)
        toolbar.addSeparator()

        bold_btn = QAction(qta.icon('fa5s.bold', color='#409eff'), "加粗", self)
        bold_btn.triggered.connect(self.toggle_bold)
        toolbar.addAction(bold_btn)

        italic_btn = QAction(qta.icon('fa5s.italic', color='#409eff'), "斜体", self)
        italic_btn.triggered.connect(self.toggle_italic)
        toolbar.addAction(italic_btn)

        underline_btn = QAction(qta.icon('fa5s.underline', color='#409eff'), "下划线", self)
        underline_btn.triggered.connect(self.toggle_underline)
        toolbar.addAction(underline_btn)

        color_btn = QAction(qta.icon('fa5s.palette', color='#409eff'), "文字颜色", self)
        color_btn.triggered.connect(self.change_text_color)
        toolbar.addAction(color_btn)
        toolbar.addSeparator()

        left_btn = QAction(qta.icon('fa5s.align-left', color='#409eff'), "左对齐", self)
        left_btn.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignLeft))
        toolbar.addAction(left_btn)

        center_btn = QAction(qta.icon('fa5s.align-center', color='#409eff'), "居中", self)
        center_btn.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignCenter))
        toolbar.addAction(center_btn)

        right_btn = QAction(qta.icon('fa5s.align-right', color='#409eff'), "右对齐", self)
        right_btn.triggered.connect(lambda: self.text_edit.setAlignment(Qt.AlignRight))
        toolbar.addAction(right_btn)
        toolbar.addSeparator()

        find_btn = QAction(qta.icon('fa5s.search', color='#409eff'), "查找替换", self)
        find_btn.triggered.connect(self.show_find_replace)
        toolbar.addAction(find_btn)
        main_layout.addWidget(toolbar)

        # 文本编辑区
        self.text_edit = QTextEdit()
        self.text_edit.setPlaceholderText(
            "在此输入文档内容...\n\n"
            "支持格式：粗体、斜体、下划线、文字颜色、对齐\n"
            "插入图片：Ctrl+G  查找替换：Ctrl+F\n"
            "右键菜单：符号处理（核心功能）"
        )
        self.text_edit.setFont(QFont(THEME_FONT, 11))
        self.text_edit.setAcceptRichText(True)
        self.text_edit.textChanged.connect(self.on_content_changed)
        self.text_edit.setContextMenuPolicy(Qt.CustomContextMenu)
        self.text_edit.customContextMenuRequested.connect(self.show_context_menu)
        main_layout.addWidget(self.text_edit)

        # 状态栏
        self.status_bar = QStatusBar()
        self.file_info_label = QLabel("未打开文件")
        self.status_bar.addWidget(self.file_info_label)
        self.word_count_label = QLabel("字数: 0")
        self.status_bar.addPermanentWidget(self.word_count_label)
        self.modified_indicator = QLabel("")
        self.status_bar.addPermanentWidget(self.modified_indicator)
        main_layout.addWidget(self.status_bar)

    # ---------- 右键菜单 ----------
    def show_context_menu(self, pos):
        if not self.is_pro:
            menu = QMenu()
            copy_action = menu.addAction("复制")
            copy_action.triggered.connect(self.text_edit.copy)
            menu.addSeparator()
            upgrade_action = menu.addAction("⚠️ 升级 PRO 版解锁全部功能")
            upgrade_action.triggered.connect(self._show_upgrade_dialog)
            menu.exec_(self.text_edit.mapToGlobal(pos))
        else:
            helper = WordContextMenuHelper(self.text_edit, self)
            menu = helper.create_context_menu(pos)
            menu.exec_(self.text_edit.mapToGlobal(pos))

    def _show_upgrade_dialog(self):
        """显示升级对话框"""
        if self.main_window and hasattr(self.main_window, '_show_activation_dialog'):
            self.main_window._show_activation_dialog()
        elif self.main_window and hasattr(self.main_window, 'show_toast'):
            self.main_window.show_toast("请联系客服获取 PRO 版激活码")

    # 在 WordEditDialog 类中添加以下方法（建议放在 create_symbol_menu 方法附近）

    def apply_symbol_to_selection(self, sym_key, position):
        """将当前选中文本应用符号处理"""
        if not self.is_pro:
            return
        cursor = self.text_edit.textCursor()
        if not cursor.hasSelection():
            if self.main_window:
                self.main_window.show_toast("请先选中要处理的文本")
            return
        selected = cursor.selectedText()
        if not selected:
            return

        cfg = self.get_symbol_config()
        if sym_key not in cfg:
            return

        left, right = cfg[sym_key][0], cfg[sym_key][1] if len(cfg[sym_key]) >= 2 else cfg[sym_key][0]
        result = self._apply_symbol_to_text(selected, left, right, position, self.split_mode)
        cursor.insertText(result)
        self.modified = True
        if self.main_window:
            self.main_window.show_toast(f"已应用符号：{cfg[sym_key][3]}")

    def create_symbol_menu(self):
        """构建符号菜单：直接列出所有符号，不嵌套‘默认符号’子菜单"""
        cfg = self.get_symbol_config()
        if not cfg:
            return None

        symbol_menu = QMenu("符号处理", self)

        # 分离默认符号和自定义符号
        default_syms = {k: v for k, v in cfg.items() if not k.startswith("custom_")}
        custom_syms = {k: v for k, v in cfg.items() if k.startswith("custom_")}

        # 添加默认符号（直接添加到根菜单）
        for sym_id, (half_left, half_right, full, name) in default_syms.items():
            # 子菜单显示符号和名称，如：“”双引号
            sub_menu = QMenu(f"{half_left}{half_right} {name}", symbol_menu)
            # 添加四种操作
            actions = [
                ("头部插入（左）", "head", half_left),
                ("尾部插入（右）", "tail", half_right),
                ("两端包裹", "both", f"{half_left}文本{half_right}"),
                ("逐字符分隔", "split", f"{half_left}文{half_right}字{half_left}分{half_right}隔"),
            ]
            for action_name, pos, preview in actions:
                act = QAction(f"{action_name}：{preview}", sub_menu)
                act.triggered.connect(lambda checked, s=sym_id, p=pos: self.apply_symbol_to_selection(s, p))
                sub_menu.addAction(act)
            symbol_menu.addMenu(sub_menu)

        # 如果有自定义符号，添加分隔符后再添加
        if custom_syms and self.is_pro:
            symbol_menu.addSeparator()
            for sym_id, (half_left, half_right, full, name) in custom_syms.items():
                sub_menu = QMenu(f"⭐ {half_left}{half_right} {name}", symbol_menu)
                actions = [
                    ("头部插入（左）", "head", half_left),
                    ("尾部插入（右）", "tail", half_right),
                    ("两端包裹", "both", f"{half_left}文本{half_right}"),
                    ("逐字符分隔", "split", f"{half_left}文{half_right}字{half_left}分{half_right}隔"),
                ]
                for action_name, pos, preview in actions:
                    act = QAction(f"{action_name}：{preview}", sub_menu)
                    act.triggered.connect(lambda checked, s=sym_id, p=pos: self.apply_symbol_to_selection(s, p))
                    sub_menu.addAction(act)
                symbol_menu.addMenu(sub_menu)

        return symbol_menu

    def get_symbol_config(self):
        """从主窗口获取符号配置（完全依赖主窗口，不设备用）"""
        if self.main_window and hasattr(self.main_window, 'get_symbol_config'):
            cfg = self.main_window.get_symbol_config()
            if cfg:          # 确保配置非空
                return cfg
        # 如果主窗口没有有效配置，返回空字典（右键菜单将不显示符号处理）
        return {}

    def _apply_symbol_to_text(self, text, left, right, position, split_mode):
        # 保留原始文本（不去除首尾空格）
        if not text:
            return text
        if position == "head":
            return left + text
        elif position == "tail":
            return text + right
        elif position == "both":
            return left + text + right
        elif position == "split":
            # 逐字符处理，保留空格和换行符作为普通字符
            chars = list(text)
            if split_mode == 0:
                return left.join(chars)
            elif split_mode == 1:
                return "".join([f"{left}{c}{right}" for c in chars])
            elif split_mode == 2:
                return ",".join([f"{left}{c}{right}" for c in chars])
        return text

    # ---------- 文件操作（优化进度和光标）----------
    def open_file(self):
        if not self.is_pro:
            QMessageBox.warning(self, "权限不足", "打开文件为 PRO 版专属")
            return
        file_path, _ = QFileDialog.getOpenFileName(self, "打开 Word 文档", "", "Word文件 (*.docx)")
        if not file_path:
            return

        # 使用进度对话框，避免鼠标转圈
        progress = QProgressDialog("正在打开文档...", "取消", 0, 100, self)
        progress.setWindowTitle("请稍候")
        progress.setMinimumDuration(0)
        progress.setValue(10)

        try:
            from docx import Document
            doc = Document(file_path)

            progress.setValue(30)
            QApplication.processEvents()

            # 简化读取：提取纯文本 + 基本格式
            html_parts = []
            for para in doc.paragraphs:
                text = para.text
                if text.strip():
                    html_parts.append(f"<p>{text}</p>")
            progress.setValue(80)
            QApplication.processEvents()

            self.text_edit.setHtml("<div>" + "".join(html_parts) + "</div>")
            self.current_path = file_path
            self.modified = False
            self._add_recent_file(file_path)
            self.setWindowTitle(f"文本处理 - {os.path.basename(file_path)}")
            self.update_status()
            progress.setValue(100)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"打开文件失败：{str(e)}")
        finally:
            progress.close()

    def _save_to_path(self, file_path):
        """静默保存（无等待光标、无弹窗）"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            block = self.text_edit.document().begin()
            while block.isValid():
                para = doc.add_paragraph()
                layout = block.blockFormat()
                alignment = layout.alignment()
                if alignment == Qt.AlignLeft:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alignment == Qt.AlignCenter:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == Qt.AlignRight:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == Qt.AlignJustify:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                it = block.begin()
                while not it.atEnd():
                    fragment = it.fragment()
                    if fragment.isValid():
                        text = fragment.text()
                        if text:
                            run = para.add_run(text)
                            font = fragment.charFormat().font()
                            if font.family():
                                run.font.name = font.family()
                            run.font.size = Pt(font.pointSize())
                            run.bold = font.bold()
                            run.italic = font.italic()
                            run.underline = font.underline()
                            fg = fragment.charFormat().foreground().color()
                            if fg.isValid():
                                run.font.color.rgb = RGBColor(fg.red(), fg.green(), fg.blue())
                    it += 1
                block = block.next()
                QApplication.processEvents()   # 保持界面响应

            doc.save(file_path)
            self.modified = False
            self.update_status()
            self.status_bar.showMessage("保存成功", 2000)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败：{str(e)}")

    def save_file(self):
        if not self.is_pro:
            QMessageBox.warning(self, "权限不足", "保存文件为 PRO 版专属")
            return
        if not self.current_path:
            self.save_file_as()
            return
        if not self.current_path.lower().endswith('.docx'):
            QMessageBox.warning(self, "提示", "保存路径必须是 .docx 文件")
            return
        try:
            self._save_to_path(self.current_path)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败：{str(e)}")

    def save_file_as(self):
        if not self.is_pro:
            QMessageBox.warning(self, "权限不足", "另存为 PRO 版专属")
            return
        default_name = "未命名.docx"
        if self.current_path:
            default_name = os.path.basename(self.current_path)
        file_path, _ = QFileDialog.getSaveFileName(self, "另存为", default_name, "Word文件 (*.docx)")
        if not file_path:
            return
        if not file_path.lower().endswith('.docx'):
            file_path += '.docx'
        try:
            self._save_to_path(file_path)
            self.current_path = file_path
            self.setWindowTitle(f"文本处理 - {os.path.basename(file_path)}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败：{str(e)}")

    def new_document(self):
        if not self.is_pro:
            QMessageBox.warning(self, "权限不足", "新建文档为 PRO 版专属")
            return
        if self.modified:
            reply = QMessageBox.question(self, "提示", "当前文档有未保存的修改，是否保存？",
                                         QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            if reply == QMessageBox.Save:
                self.save_file()
                if self.modified:
                    return
            elif reply == QMessageBox.Cancel:
                return
        self.text_edit.clear()
        self.current_path = None
        self.modified = False
        self.update_status()
        self.setWindowTitle("文本处理 - 新建文档")

    # ---------- 其他功能（插入图片、格式化、查找替换等）----------
    def insert_image(self):
        if not self.is_pro:
            QMessageBox.warning(self, "权限不足", "插入图片功能为 PRO 版专属")
            return
        file_path, _ = QFileDialog.getOpenFileName(self, "选择图片", "", "图片文件 (*.png *.jpg *.jpeg *.bmp *.gif)")
        if not file_path:
            return
        try:
            with open(file_path, 'rb') as f:
                img_data = base64.b64encode(f.read()).decode('utf-8')
            img_html = f'<img src="data:image/png;base64,{img_data}" style="max-width: 100%;" />'
            self.text_edit.textCursor().insertHtml(img_html)
            self.modified = True
            self.status_bar.showMessage("图片已插入", 2000)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"插入图片失败：{str(e)}")

    def toggle_bold(self):
        fmt = self.text_edit.currentCharFormat()
        fmt.setFontWeight(QFont.Bold if fmt.fontWeight() != QFont.Bold else QFont.Normal)
        self.text_edit.setCurrentCharFormat(fmt)
        self.modified = True

    def toggle_italic(self):
        fmt = self.text_edit.currentCharFormat()
        fmt.setFontItalic(not fmt.fontItalic())
        self.text_edit.setCurrentCharFormat(fmt)
        self.modified = True

    def toggle_underline(self):
        fmt = self.text_edit.currentCharFormat()
        fmt.setFontUnderline(not fmt.fontUnderline())
        self.text_edit.setCurrentCharFormat(fmt)
        self.modified = True

    def change_text_color(self):
        current_fmt = self.text_edit.currentCharFormat()
        initial_color = current_fmt.foreground().color() if current_fmt.foreground().color().isValid() else QColor("#000000")
        color_dlg = ColorPickerDialog(self, initial_color)
        if color_dlg.exec():
            color = color_dlg.get_color()
            fmt = self.text_edit.currentCharFormat()
            fmt.setForeground(color)
            self.text_edit.setCurrentCharFormat(fmt)
            self.modified = True

    def show_find_replace(self):
        dlg = FindReplaceDialog(self)
        dlg.exec()

    def on_content_changed(self):
        if self.current_path and not self.modified:
            self.modified = True
            self.update_status()
        text = self.text_edit.toPlainText()
        word_count = len(text.replace(" ", "").replace("\n", ""))
        self.word_count_label.setText(f"字数: {word_count}")

    def update_status(self):
        if self.current_path:
            self.file_info_label.setText(os.path.basename(self.current_path))
        else:
            self.file_info_label.setText("未打开文件")
        if self.modified:
            self.modified_indicator.setText("● 已修改")
            self.modified_indicator.setStyleSheet("color: #f56c6c; font-weight: bold;")
        else:
            self.modified_indicator.setText("")

    def setup_shortcuts(self):
        from PySide6.QtGui import QShortcut, QKeySequence
        self._shortcuts = []
        for key, callback in [
            (QKeySequence.Find, self.show_find_replace),
            (QKeySequence.Bold, self.toggle_bold),
            (QKeySequence.Italic, self.toggle_italic),
            (QKeySequence.Underline, self.toggle_underline),
            (QKeySequence("Ctrl+G"), self.insert_image),
        ]:
            shortcut = QShortcut(key, self)
            shortcut.activated.connect(callback)
            self._shortcuts.append(shortcut)

    def _set_readonly_mode(self):
        self.text_edit.setReadOnly(True)
        self.text_edit.setPlaceholderText(
            "【免费版】仅可查看文档内容，如需编辑请升级到 PRO 版。\n"
            "💡 升级后可使用加粗、斜体、插入图片、保存文件等完整功能。"
        )
    
    def enable_pro_features(self):
        """激活后启用所有PRO版功能"""
        self.is_pro = True
        self.text_edit.setReadOnly(False)
        self.text_edit.setPlaceholderText("在此输入文档内容...")
        
        if hasattr(self, 'toolbar'):
            for child in self.toolbar.findChildren(QPushButton):
                child.setEnabled(True)
        
        for action in self.findChildren(QAction):
            text = action.text()
            if text in ["新建", "保存", "另存为..."]:
                action.setEnabled(True)
    
    def set_language(self, lang):
        """设置符号语言（避免报错）"""
        pass

    def apply_theme(self):
        if not self.theme_manager:
            return
        theme = self.theme_manager.get_theme()
        primary = theme["primary"]
        bg = theme["bg"]
        card = theme["card"]
        border = theme["border"]
        text = theme["text"]
        hover_bg = theme.get("hover_bg", f"{primary}15")
        self.setStyleSheet(f"""
            QTextEdit {{
                background-color: {card};
                border: 1px solid {border};
                border-radius: 4px;
                padding: 10px;
                font-family: "{THEME_FONT}", "Segoe UI", Arial;
                font-size: 11pt;
                color: {text};
            }}
            QTextEdit:focus {{ border-color: {primary}; }}
            QStatusBar {{
                background-color: {bg};
                border-top: 1px solid {border};
                padding: 4px;
                color: {text};
            }}
            QMenuBar {{
                background-color: {bg};
                border-bottom: 1px solid {border};
                color: {text};
            }}
            QMenuBar::item:selected {{
                background-color: {primary};
                color: white;
            }}
            QToolBar {{
                background-color: {card};
                border-bottom: 1px solid {border};
                padding: 4px;
                spacing: 8px;
            }}
            QToolButton {{
                border: none;
                padding: 6px;
                border-radius: 4px;
                color: {text};
            }}
            QToolButton:hover {{
                background-color: {hover_bg};
                color: {primary};
            }}
            QToolButton:pressed {{
                background-color: {border};
            }}
            QMenu {{
                background-color: {card};
                border: 1px solid {border};
                border-radius: 8px;
                padding: 6px;
                color: {text};
            }}
            QMenu::item:selected {{
                background-color: {hover_bg};
                color: {primary};
            }}
        """)

    def closeEvent(self, event):
        if self.modified and self.is_pro:
            reply = QMessageBox.question(self, "提示", "文档已修改，是否保存？",
                                         QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            if reply == QMessageBox.Save:
                self.save_file()
                if self.modified:
                    event.ignore()
                    return
            elif reply == QMessageBox.Cancel:
                event.ignore()
                return
        event.accept()

