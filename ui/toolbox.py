# ui/toolbox.py
import os
import sys
import shutil
import zipfile
import json
import traceback
from datetime import datetime
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QGridLayout, QFrame, QLabel, QDialog,
    QListWidget, QPushButton, QLineEdit, QSpinBox, QComboBox,
    QFileDialog, QMessageBox, QInputDialog, QApplication, QTabWidget, QHBoxLayout,
    QSizePolicy, QTextEdit, QScrollArea, QProgressBar, QTextBrowser, QCheckBox, QAbstractItemView,
    QTableWidget, QTableWidgetItem, QHeaderView
)
from PySide6.QtCore import Qt, QTimer, Signal, Slot, QSize, QThread, QMimeData, QRect
from PySide6.QtGui import QPixmap, QImage, QFont, QIcon, QDragEnterEvent, QDropEvent, QPainter, QPen, QColor, QCursor

import qtawesome as qta

from utils.logger import get_logger

logger = get_logger()



ICON_SIZE_SMALL = QSize(16, 16)
ICON_SIZE_MEDIUM = QSize(20, 20)
ICON_SIZE_LARGE = QSize(24, 24)


class DragDropListWidget(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setDragDropMode(QAbstractItemView.DropOnly)
        self.file_list = []

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragMoveEvent(event)

    def dropEvent(self, event: QDropEvent):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            for url in urls:
                file_path = url.toLocalFile()
                if os.path.isfile(file_path):
                    if file_path not in self.file_list:
                        self.file_list.append(file_path)
                        self.addItem(file_path)
            event.acceptProposedAction()
        else:
            super().dropEvent(event)

    def clear_files(self):
        self.file_list = []
        self.clear()

    def get_files(self):
        return self.file_list


class ConvertThread(QThread):
    progress = Signal(int)
    log = Signal(str)
    finished = Signal(bool, str)

    def __init__(self, task_func, *args, **kwargs):
        super().__init__()
        self.task_func = task_func
        self.args = args
        self.kwargs = kwargs
        self.is_cancelled = False

    def cancel(self):
        self.is_cancelled = True

    def run(self):
        try:
            if self.is_cancelled:
                self.log.emit("任务已取消")
                self.finished.emit(False, "已取消")
                return
                
            # 将线程自身传递给任务函数，这样任务可以实时检查 is_cancelled
            result = self.task_func(self.progress, self.log, self, *self.args, **self.kwargs)
            if self.is_cancelled:
                self.finished.emit(False, "已取消")
                return
            self.finished.emit(True, result or "转换完成")
        except Exception as e:
            if self.is_cancelled:
                self.finished.emit(False, "已取消")
            else:
                import traceback
                self.finished.emit(False, f"错误：{str(e)}\n{traceback.format_exc()}")


class OcrThread(QThread):
    progress = Signal(int)
    log = Signal(str)
    finished = Signal(bool, str)

    def __init__(self, image_path, tesseract_cmd, tessdata_dir, *args, **kwargs):
        super().__init__()
        self.image_path = image_path
        self.tesseract_cmd = tesseract_cmd
        self.tessdata_dir = tessdata_dir
        self.args = args
        self.kwargs = kwargs
        self.is_cancelled = False
        self.current_process = None

    def cancel(self):
        self.is_cancelled = True
        # 如果有正在运行的进程，强制终止它
        if self.current_process:
            try:
                self.current_process.kill()
            except:
                pass

    def run(self):
        try:
            import pytesseract
            from PIL import Image
            import subprocess
            import tempfile
            from PIL import ImageEnhance
            import time

            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd

            if self.is_cancelled:
                self.log.emit("识别已取消")
                self.finished.emit(False, "已取消")
                return

            self.log.emit("正在预处理图片...")
            self.progress.emit(10)
            
            img = Image.open(self.image_path)
            if img.mode != 'L':
                img = img.convert('L')
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.5)

            psm_options = ['3', '6', '4', '1']
            recognized_text = ""
            best_text = ""

            for idx, psm in enumerate(psm_options):
                if self.is_cancelled:
                    self.log.emit("识别已取消")
                    self.finished.emit(False, "已取消")
                    return
                
                self.log.emit(f"正在识别 (模式 {psm})...")
                progress_base = 15 + (idx * 20)
                self.progress.emit(progress_base)
                
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    img.save(tmp_img.name, 'PNG')
                    tmp_img_path = tmp_img.name
                with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp_out:
                    tmp_out_path = tmp_out.name[:-4]
                try:
                    cmd = [self.tesseract_cmd]
                    if self.tessdata_dir and os.path.exists(self.tessdata_dir):
                        cmd.extend(['--tessdata-dir', self.tessdata_dir])
                    cmd.extend([tmp_img_path, tmp_out_path, '-l', 'chi_sim+eng', '--psm', psm])
                    
                    # 隐藏tesseract终端窗口（Windows）
                    startupinfo = None
                    if sys.platform.startswith('win'):
                        startupinfo = subprocess.STARTUPINFO()
                        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                        startupinfo.wShowWindow = subprocess.SW_HIDE
                    
                    # 使用Popen而不是run，以便支持取消
                    self.current_process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, 
                                                          startupinfo=startupinfo)
                    
                    # 轮询等待完成，同时检查取消标志
                    while self.current_process.poll() is None:
                        if self.is_cancelled:
                            try:
                                self.current_process.kill()
                            except:
                                pass
                            self.log.emit("识别已取消")
                            self.finished.emit(False, "已取消")
                            return
                        time.sleep(0.1)
                    
                    self.current_process = None
                    output_txt_path = tmp_out_path + '.txt'
                    if os.path.exists(output_txt_path):
                        try:
                            with open(output_txt_path, 'r', encoding='utf-8') as f:
                                recognized_text = f.read()
                        except UnicodeDecodeError:
                            try:
                                with open(output_txt_path, 'r', encoding='gbk') as f:
                                    recognized_text = f.read()
                            except:
                                with open(output_txt_path, 'r', encoding='gb18030', errors='ignore') as f:
                                    recognized_text = f.read()
                        if len(recognized_text.strip()) > len(best_text.strip()):
                            best_text = recognized_text
                        self.progress.emit(progress_base + 15)
                finally:
                    self.current_process = None
                    if os.path.exists(tmp_img_path):
                        try:
                            os.unlink(tmp_img_path)
                        except:
                            pass
                    if os.path.exists(tmp_out_path + '.txt'):
                        try:
                            os.unlink(tmp_out_path + '.txt')
                        except:
                            pass

            final_text = best_text if best_text else recognized_text
            self.progress.emit(100)
            if final_text.strip():
                self.log.emit("识别完成！")
                self.finished.emit(True, final_text.strip())
            else:
                self.log.emit("未识别到文字")
                self.finished.emit(True, "未识别到文字")
        except ImportError as e:
            self.finished.emit(False, f"缺少依赖库: {str(e)}")
        except Exception as e:
            self.finished.emit(False, f"识别失败: {str(e)}")


class AudioThread(QThread):
    progress = Signal(int)
    log = Signal(str)
    finished = Signal(bool, str)

    def __init__(self, file_list, output_folder, output_format, ffmpeg_path, *args, **kwargs):
        super().__init__()
        self.file_list = file_list
        self.output_folder = output_folder
        self.output_format = output_format
        self.ffmpeg_path = ffmpeg_path
        self.args = args
        self.kwargs = kwargs
        self.is_cancelled = False
        self.current_output_file = None  # 当前正在处理的输出文件
        self.process = None  # 当前子进程

    def cancel(self):
        self.is_cancelled = True
        if self.process:
            try:
                self.process.kill()
            except:
                pass

    def run(self):
        import subprocess
        total = len(self.file_list)
        success_count = 0
        error_count = 0
        os.makedirs(self.output_folder, exist_ok=True)
        
        for idx, input_file in enumerate(self.file_list):
            if self.is_cancelled:
                self.log.emit("转换已取消")
                break
            base_name = os.path.splitext(os.path.basename(input_file))[0]
            self.current_output_file = os.path.join(self.output_folder, f"{base_name}.{self.output_format}")
            progress = int((idx / total) * 100)
            self.progress.emit(progress)
            self.log.emit(f"正在转换 ({idx+1}/{total}): {base_name}...")
            
            try:
                cmd = [
                    self.ffmpeg_path, '-i', input_file,
                    '-y', '-hide_banner',
                    '-threads', '4'
                ]
                if self.output_format == 'ogg':
                    cmd.extend(['-c:a', 'libvorbis', '-q:a', '6'])
                elif self.output_format == 'mp3':
                    cmd.extend(['-c:a', 'libmp3lame', '-q:a', '2'])
                elif self.output_format == 'flac':
                    cmd.extend(['-c:a', 'flac', '-compression_level', '5'])
                elif self.output_format == 'wav':
                    cmd.extend(['-c:a', 'pcm_s16le'])
                elif self.output_format == 'aac' or self.output_format == 'm4a':
                    cmd.extend(['-c:a', 'aac', '-b:a', '256k'])
                elif self.output_format == 'opus':
                    cmd.extend(['-c:a', 'libopus', '-b:a', '128k'])
                elif self.output_format == 'ape':
                    cmd.extend(['-c:a', 'ape', '-q:a', '8'])
                elif self.output_format == 'wma':
                    cmd.extend(['-c:a', 'wmav2', '-b:a', '192k'])
                cmd.append('-progress')
                cmd.append('pipe:1')
                cmd.append(self.current_output_file)
                
                self.process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                stdout, stderr = self.process.communicate(timeout=300)
                
                if self.process.returncode != 0:
                    error_msg = stderr if stderr else "未知错误"
                    if 'Invalid' in error_msg or 'Unsupported' in error_msg:
                        self.log.emit(f"✗ 格式不支持: {base_name}")
                    else:
                        self.log.emit(f"✗ 转换失败: {base_name} - {error_msg[:80]}")
                    error_count += 1
                    # 失败时删除不完整的文件
                    if os.path.exists(self.current_output_file):
                        try:
                            os.remove(self.current_output_file)
                        except:
                            pass
                else:
                    if os.path.exists(self.current_output_file) and os.path.getsize(self.current_output_file) > 0:
                        success_count += 1
                        self.log.emit(f"✓ 转换成功: {base_name}")
                    else:
                        error_count += 1
                        self.log.emit(f"✗ 文件无效: {base_name}")
                        if os.path.exists(self.current_output_file):
                            try:
                                os.remove(self.current_output_file)
                            except:
                                pass
            except subprocess.TimeoutExpired:
                if self.process:
                    self.process.kill()
                error_count += 1
                self.log.emit(f"✗ 转换超时: {base_name}")
                if os.path.exists(self.current_output_file):
                    try:
                        os.remove(self.current_output_file)
                    except:
                        pass
            except Exception as e:
                error_count += 1
                self.log.emit(f"✗ 转换出错: {base_name} - {str(e)[:80]}")
                if os.path.exists(self.current_output_file):
                    try:
                        os.remove(self.current_output_file)
                    except:
                        pass
            finally:
                self.process = None
                self.current_output_file = None
        
        self.progress.emit(100)
        if self.is_cancelled:
            self.finished.emit(False, f"转换已取消。成功: {success_count}, 失败: {error_count}")
        elif error_count > 0:
            self.finished.emit(True, f"转换完成！成功: {success_count}, 失败: {error_count}")
        else:
            self.finished.emit(True, f"全部转换成功！{success_count} 个文件")


class ConvertProgressDialog(QDialog):
    def __init__(self, title="处理中", parent=None, theme=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setMinimumSize(500, 350)
        self.setModal(True)
        self.theme = theme or {}

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)

        bg_color = self.theme.get("bg", "#ffffff")
        card_bg = self.theme.get("card", "#ffffff")
        text_color = self.theme.get("text", "#303133")
        border_color = self.theme.get("border", "#dcdfe6")
        primary_color = self.theme.get("primary", "#409eff")
        primary_dark = self.theme.get("primary_dark", "#337ecc")
        success_color = self.theme.get("success", "#67c23a")
        error_color = self.theme.get("error", "#f56c6c")
        
        # 保存颜色用于后续使用
        self.success_color = success_color
        self.error_color = error_color

        self.title_label = QLabel(title)
        self.title_label.setStyleSheet(f"font-size: 16px; font-weight: bold; color: {text_color};")
        layout.addWidget(self.title_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setStyleSheet(f"""
            QProgressBar {{
                border: 1px solid {border_color};
                border-radius: 6px;
                text-align: center;
                background-color: {bg_color};
                color: {text_color};
            }}
            QProgressBar::chunk {{
                border-radius: 5px;
                background-color: {primary_color};
            }}
        """)
        layout.addWidget(self.progress_bar)

        self.log_browser = QTextBrowser()
        self.log_browser.setMaximumHeight(200)
        self.log_browser.setStyleSheet(f"""
            QTextBrowser {{
                border: 1px solid {border_color};
                border-radius: 8px;
                padding: 8px;
                background-color: {card_bg};
                color: {text_color};
            }}
        """)
        layout.addWidget(self.log_browser)

        self.status_label = QLabel("准备中...")
        self.status_label.setStyleSheet(f"color: {text_color}; font-size: 13px; opacity: 0.7;")
        layout.addWidget(self.status_label)

        btn_style = f"""
            QPushButton {{
                padding: 8px 24px;
                background-color: {primary_color};
                color: white;
                border: none;
                border-radius: 6px;
                font-size: 14px;
            }}
            QPushButton:hover {{
                background-color: {primary_dark};
            }}
        """
        self.cancel_btn = QPushButton("取消")
        self.cancel_btn.setStyleSheet(btn_style)
        self.cancel_btn.clicked.connect(self.reject)
        layout.addWidget(self.cancel_btn)

        self.thread = None
        self.is_running = False

    def start_convert(self, task_func, *args, **kwargs):
        self.is_running = True
        self.thread = ConvertThread(task_func, *args, **kwargs)
        self.thread.progress.connect(self.update_progress)
        self.thread.log.connect(self.add_log)
        self.thread.finished.connect(self.on_finished)
        self.thread.start()

    def update_progress(self, value):
        self.progress_bar.setValue(value)
        self.status_label.setText(f"处理中... {value}%")

    def add_log(self, msg):
        self.log_browser.append(msg)
        self.log_browser.verticalScrollBar().setValue(
            self.log_browser.verticalScrollBar().maximum()
        )

    def on_finished(self, success, message):
        self.is_running = False
        self.progress_bar.setValue(100 if success else 0)
        self.status_label.setText("处理完成" if success else "处理失败")
        self.status_label.setStyleSheet(
            f"color: {self.success_color}; font-size: 13px;" if success else f"color: {self.error_color}; font-size: 13px;"
        )
        self.add_log(message)
        self.cancel_btn.setText("关闭")

    def closeEvent(self, event):
        if self.is_running and self.thread:
            # 使用标志位优雅停止线程，而不是强制终止
            self.thread.cancel()
            self.thread.wait(5000)  # 最多等待5秒
        super().closeEvent(event)


class ToolBoxWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.main_window = parent
        self.is_pro = False
        if self.main_window and hasattr(self.main_window, 'activation_manager'):
            self.is_pro = self.main_window.activation_manager.is_activated()

        self.clipboard_history = []
        self.clipboard = QApplication.clipboard()
        self.clipboard.dataChanged.connect(self._on_clipboard_change)
        # 敏感信息过滤开关，从配置读取，默认开启
        self.sensitive_filter_enabled = True
        if self.main_window and hasattr(self.main_window, 'config_manager'):
            config_value = self.main_window.config_manager.get_config("ToolBox", "sensitive_filter_enabled", "true")
            self.sensitive_filter_enabled = config_value.lower() == "true"

        self._init_ui()

        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            self.apply_theme()

    # ---------- 权限检查 ----------
    def _check_pro_or_warn(self) -> bool:
        """检查是否为 PRO 版，若不是则显示警告并返回 False"""
        if not self.is_pro:
            self.show_warning("PRO版专属功能", "此功能需要升级到 PRO 版才能使用。\n请联系客服获取激活码。")
            return False
        return True

    def show_warning(self, title, text):
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle(title)
            msg_box.setText(text)
            msg_box.setIcon(QMessageBox.Warning)
            msg_box.setStyleSheet(self.main_window.theme_manager.generate_stylesheet())
            msg_box.exec()
        else:
            QMessageBox.warning(self, title, text)

    def show_error(self, title, text):
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle(title)
            msg_box.setText(text)
            msg_box.setIcon(QMessageBox.Critical)
            msg_box.setStyleSheet(self.main_window.theme_manager.generate_stylesheet())
            msg_box.exec()
        else:
            QMessageBox.critical(self, title, text)

    def show_info(self, title, text):
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle(title)
            msg_box.setText(text)
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setStyleSheet(self.main_window.theme_manager.generate_stylesheet())
            msg_box.exec()
        else:
            QMessageBox.information(self, title, text)

    # ---------- UI 刷新与主题 ----------
    def refresh_activation_status(self):
        """刷新激活状态，激活后无限制使用"""
        if self.main_window and hasattr(self.main_window, 'activation_manager'):
            self.is_pro = self.main_window.activation_manager.is_activated()
        # 重新构建UI以更新状态
        if hasattr(self, 'grid_layout'):
            # 清除现有内容
            while self.grid_layout.count():
                item = self.grid_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            # 重新添加工具卡片
            self.tool_cards = []
            tools = [
                ("剪贴板历史", "保存复制记录，一键复用", self.open_clipboard_history),
                ("批量文件处理", "批量改名/内容替换/编号", self.open_batch_processing),
                ("图片工具", "压缩、格式转换、ICO图标生成", self.open_image_tools),
                ("压缩解压", "打包解压、批量ZIP", self.open_zip_tools),
                ("批量文字提取", "Word/Excel/PDF/TXT/图片提取文字", self.open_batch_text_extract),
                ("图片OCR", "图片文字识别", self.open_image_ocr),
                ("文本格式化", "JSON/SQL/Markdown一键排版", self.open_text_format),
                ("格式互转", "Word/Excel/PDF/HTML互转", self.open_format_convert),
                ("音频转换", "音频转MP3/WAV/FLAC/AAC/OGG/opus", self.open_audio_convert),
            ]
            tool_icons = [
                'fa5s.clipboard',
                'fa5s.folder-open',
                'fa5s.image',
                'fa5s.file-archive',
                'fa5s.file-alt',
                'fa5s.magic',
                'fa5s.code',
                'fa5s.exchange-alt',
                'fa5s.music'
            ]
            for i, (title, desc, callback) in enumerate(tools):
                row = i // 3
                col = i % 3
                icon_name = tool_icons[i]
                card = self._create_tool_card(icon_name, title, desc, callback)
                self.tool_cards.append(card)
                self.grid_layout.addWidget(card, row, col)

    def apply_theme(self):
        if not self.main_window or not hasattr(self.main_window, 'theme_manager'):
            return
        theme = self.main_window.theme_manager.get_theme()
        card_bg = theme["card"]
        text = theme["text"]
        text_secondary = theme.get("text_secondary", "#6b7280")
        border = theme["border"]
        primary = theme["primary"]
        primary_light = theme.get("primary_light", primary)
        bg = theme["bg"]

        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip('#')
            if len(hex_color) == 6:
                return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            return (64, 158, 255)

        r, g, b = hex_to_rgb(primary)
        rgba_15 = f"rgba({r}, {g}, {b}, 0.15)"
        rgba_20 = f"rgba({r}, {g}, {b}, 0.20)"

        self.setStyleSheet(f"""
            QWidget {{
                background-color: {bg};
            }}
            QFrame#toolCard {{
                background-color: {card_bg};
                border: 1px solid {border};
                border-radius: 12px;
                padding: 0;
                background-image: none;
            }}
            QFrame#toolCard:hover {{
                background-color: {rgba_15};
                border: 1px solid {primary};
                background-image: none;
            }}
            QFrame#toolCard:pressed {{
                background-color: {primary_light};
                border: 1px solid {primary};
                background-image: none;
            }}
            QFrame#toolCard > QWidget, QFrame#toolCard > QLayout > QWidget {{
                background-color: transparent;
                background-image: none;
            }}
            QLabel#toolCardTitle {{
                font-size: 15px;
                font-weight: 600;
                color: {text};
                background-color: transparent;
            }}
            QLabel#toolCardDesc {{
                font-size: 12px;
                color: {text_secondary};
                background-color: transparent;
            }}
            QLabel#toolCardLock {{
                font-size: 11px;
                color: #f56c6c;
                background-color: transparent;
                padding-top: 4px;
            }}
            QLabel {{
                color: {text};
                background-color: transparent;
            }}
            QLabel#toolboxTitle {{
                color: {text};
                font-size: 18px;
                font-weight: bold;
                background: transparent;
            }}
            QListWidget {{
                background-color: {card_bg};
                border: 1px solid {border};
                border-radius: 6px;
            }}
            QListWidget::item:hover {{
                background-color: {rgba_20};
            }}
            QPushButton {{
                background-color: {primary};
                color: white;
                border: none;
                border-radius: 6px;
                padding: 6px 12px;
            }}
            QPushButton:hover {{
                background-color: {primary_light};
            }}
            QPushButton:disabled {{
                background-color: {border};
                color: {text};
            }}
            QLineEdit, QComboBox, QSpinBox, QTextEdit {{
                background-color: {card_bg};
                border: 1px solid {border};
                border-radius: 4px;
                padding: 4px;
                color: {text};
            }}
            QLineEdit:focus, QComboBox:focus, QSpinBox:focus, QTextEdit:focus {{
                border: 1px solid {primary};
            }}
            QTabWidget::pane {{
                background-color: {card_bg};
                border: 1px solid {border};
            }}
            QTabBar::tab {{
                background-color: {card_bg};
                color: {text};
                padding: 6px 12px;
            }}
            QTabBar::tab:selected {{
                background-color: {primary};
                color: white;
            }}
            QTabBar::tab:hover:!selected {{
                background-color: {rgba_20};
            }}
            QScrollArea {{
                background-color: {card_bg};
                border: 1px solid {border};
            }}
            QDialog {{
                background-color: {card_bg};
            }}
        """)

        for card in self.tool_cards:
            for child in card.findChildren(QLabel):
                if child.property("icon_name"):
                    icon_name = child.property("icon_name")
                    child.setPixmap(qta.icon(icon_name, color=primary).pixmap(48, 48))

        for child in self.findChildren(QLabel):
            if child.property("title_icon"):
                child.setPixmap(qta.icon('fa5s.toolbox', color=primary).pixmap(ICON_SIZE_LARGE))

    def _init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_layout.setSpacing(12)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        title_widget = QWidget()
        title_layout = QHBoxLayout(title_widget)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(8)

        title_icon = QLabel()
        title_icon.setProperty("title_icon", True)
        title_icon.setObjectName("toolboxTitleIcon")
        title_icon.setPixmap(qta.icon('fa5s.toolbox', color='#374151').pixmap(ICON_SIZE_LARGE))
        title_layout.addWidget(title_icon)

        title_label = QLabel("工具箱")
        title_label.setObjectName("toolboxTitle")
        title_label.setStyleSheet("color: #374151; font-size: 18px; font-weight: bold; background: transparent;")
        title_layout.addWidget(title_label)
        title_layout.addStretch()

        main_layout.addWidget(title_widget)

        self.grid_layout = QGridLayout()
        self.grid_layout.setSpacing(12)
        self.grid_layout.setContentsMargins(0, 0, 0, 0)

        self.tool_cards = []
        tools = [
            ("剪贴板历史", "保存复制记录，一键复用", self.open_clipboard_history),
            ("批量文件处理", "批量改名/内容替换/编号", self.open_batch_processing),
            ("图片工具", "压缩、格式转换、ICO图标生成", self.open_image_tools),
            ("压缩解压", "打包解压、批量ZIP", self.open_zip_tools),
            ("批量文字提取", "Word/Excel/PDF/TXT/图片提取文字", self.open_batch_text_extract),
            ("图片OCR", "图片文字识别", self.open_image_ocr),
            ("文本格式化", "JSON/SQL/Markdown一键排版", self.open_text_format),
            ("格式互转", "Word/Excel/PDF/HTML互转", self.open_format_convert),
            ("音频转换", "音频转MP3/WAV/FLAC/AAC/OGG/opus", self.open_audio_convert),
        ]

        tool_icons = [
            'fa5s.clipboard',
            'fa5s.folder-open',
            'fa5s.image',
            'fa5s.file-archive',
            'fa5s.file-alt',
            'fa5s.magic',
            'fa5s.code',
            'fa5s.exchange-alt',
            'fa5s.music'
        ]

        for i, (title, desc, callback) in enumerate(tools):
            row = i // 3
            col = i % 3
            icon_name = tool_icons[i]
            card = self._create_tool_card(icon_name, title, desc, callback)
            self.tool_cards.append(card)
            self.grid_layout.addWidget(card, row, col)

        main_layout.addLayout(self.grid_layout)

        for row in range(3):
            self.grid_layout.setRowStretch(row, 1)
        for col in range(3):
            self.grid_layout.setColumnStretch(col, 1)

    def _create_tool_card(self, icon_name, title, desc, callback):
        card = QFrame()
        card.setProperty("class", "ToolCard")
        card.setObjectName("toolCard")
        
        layout = QVBoxLayout(card)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setStretch(0, 1)
        layout.setStretch(1, 0)
        layout.setStretch(2, 0)

        # 图标区域
        icon_container = QWidget()
        icon_layout = QHBoxLayout(icon_container)
        icon_layout.setContentsMargins(0, 0, 0, 0)
        icon_layout.setAlignment(Qt.AlignCenter)
        
        icon_label = QLabel()
        icon_label.setProperty("icon_name", icon_name)
        icon_label.setPixmap(qta.icon(icon_name, color='#409eff').pixmap(48, 48))
        icon_label.setAlignment(Qt.AlignCenter)
        icon_layout.addWidget(icon_label)
        layout.addWidget(icon_container)

        # 标题
        title_label = QLabel(title)
        title_label.setObjectName("toolCardTitle")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)

        # 描述
        desc_label = QLabel(desc)
        desc_label.setObjectName("toolCardDesc")
        desc_label.setAlignment(Qt.AlignCenter)
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)

        # 免费版：显示锁定提示，并修改点击事件为警告
        if not self.is_pro:
            lock_label = QLabel("🔒 PRO版专属")
            lock_label.setObjectName("toolCardLock")
            lock_label.setAlignment(Qt.AlignCenter)
            layout.addWidget(lock_label)
            # 免费版下点击弹出警告，不执行回调
            def locked_click(event):
                self._check_pro_or_warn()
            card.mousePressEvent = locked_click
        else:
            card.setCursor(Qt.PointingHandCursor)
            card.setToolTip(f"打开 {title} 工具")
            def handle_click(event):
                callback()
            card.mousePressEvent = handle_click

        return card

    def show_toast(self, message):
        if self.main_window and hasattr(self.main_window, 'show_toast'):
            self.main_window.show_toast(message)
        else:
            toast = QLabel(message, self)
            toast.setWindowFlags(Qt.ToolTip | Qt.FramelessWindowHint)
            toast.setStyleSheet("background-color: rgba(0,0,0,0.85); color: white; padding: 10px; border-radius: 8px;")
            toast.show()
            QTimer.singleShot(2000, toast.close)

    # ================== 1. 剪贴板历史 ==================
    def open_clipboard_history(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("剪贴板历史")
        dialog.setMinimumSize(500, 400)
        layout = QVBoxLayout(dialog)

        top_layout = QHBoxLayout()
        hint_label = QLabel("💡 双击列表项即可复制到剪贴板")
        hint_label.setStyleSheet("color: #6b7280; font-size: 12px;")
        top_layout.addWidget(hint_label)
        
        # 敏感信息过滤开关，从配置读取
        self.sensitive_filter_checkbox = QCheckBox("启用敏感信息过滤")
        # 每次打开时都从配置读取最新状态
        if self.main_window and hasattr(self.main_window, 'config_manager'):
            config_value = self.main_window.config_manager.get_config("ToolBox", "sensitive_filter_enabled", "true")
            self.sensitive_filter_enabled = config_value.lower() == "true"
        self.sensitive_filter_checkbox.setChecked(self.sensitive_filter_enabled)
        self.sensitive_filter_checkbox.stateChanged.connect(self._toggle_sensitive_filter)
        top_layout.addWidget(self.sensitive_filter_checkbox)
        top_layout.addStretch()
        layout.addLayout(top_layout)

        self.clipboard_list = QListWidget()
        for item in self.clipboard_history:
            self.clipboard_list.addItem(item[:50] + "..." if len(item) > 50 else item)
        self.clipboard_list.itemDoubleClicked.connect(self._copy_clipboard_item)
        layout.addWidget(self.clipboard_list)

        btn_layout = QHBoxLayout()
        clear_btn = QPushButton(" 清空历史")
        clear_btn.setIcon(qta.icon('fa5s.trash-alt', color='white'))
        clear_btn.setIconSize(ICON_SIZE_SMALL)
        clear_btn.setToolTip("清空所有剪贴板历史记录")
        clear_btn.clicked.connect(self._clear_clipboard_history)
        btn_layout.addWidget(clear_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # 保存对话框引用和关闭事件
        self.clipboard_history_dialog = dialog
        dialog.accepted.connect(self._save_clipboard_history_settings)
        dialog.rejected.connect(self._save_clipboard_history_settings)
        dialog.finished.connect(self._save_clipboard_history_settings)
        
        dialog.exec()
    
    def _save_clipboard_history_settings(self):
        """保存剪贴板历史的设置"""
        if hasattr(self, 'sensitive_filter_checkbox'):
            self.sensitive_filter_enabled = self.sensitive_filter_checkbox.isChecked()
            if self.main_window and hasattr(self.main_window, 'config_manager'):
                self.main_window.config_manager.set_config("ToolBox", "sensitive_filter_enabled", str(self.sensitive_filter_enabled).lower())

    def _toggle_sensitive_filter(self, state):
        self.sensitive_filter_enabled = state == Qt.Checked
        # 保存设置到配置
        if self.main_window and hasattr(self.main_window, 'config_manager'):
            self.main_window.config_manager.set_config("ToolBox", "sensitive_filter_enabled", str(self.sensitive_filter_enabled).lower())

    def _on_clipboard_change(self):
        text = self.clipboard.text()
        if text and text not in self.clipboard_history:
            # 敏感信息过滤（可配置）
            if self.sensitive_filter_enabled and self._is_sensitive_content(text):
                return  # 跳过记录敏感信息
            self.clipboard_history.insert(0, text)
            if len(self.clipboard_history) > 50:
                self.clipboard_history.pop()
    
    def _is_sensitive_content(self, text):
        """检查是否包含敏感信息"""
        import re
        # 常见密码/密钥模式
        sensitive_patterns = [
            # 纯数字/纯字母密码（短密码）
            r'^[0-9]{4,16}$',
            r'^[a-zA-Z]{6,16}$',
            # API密钥、令牌模式
            r'api[_-]?key',
            r'token',
            r'secret',
            # 密码提示词
            r'\bpassword\b',
            r'\bpasswd\b',
            r'\bpwd\b',
            # 信用卡/银行卡号（16-19位，带或不带分隔符）
            r'\b[0-9]{4}[\s-]?[0-9]{4}[\s-]?[0-9]{4}[\s-]?[0-9]{1,4}\b',
            # 身份证号（中国）- 18位
            r'\b[1-9][0-9]{5}(19|20)[0-9]{2}(0[1-9]|1[0-2])(0[1-9]|[12][0-9]|3[01])[0-9]{3}[0-9Xx]\b',
            # 手机号（中国）- 11位，通常以1开头
            r'\b1[3-9][0-9]{9}\b',
            # 手机号（带分隔符）
            r'\b1[3-9][\s-]?[0-9]{4}[\s-]?[0-9]{4}\b',
            # 邮箱地址
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
            # 社保卡号
            r'\b[1-9][0-9]{5}[0-9]{10}[0-9xX]\b',
            # 护照号（中国普通护照）
            r'\bE[0-9]{8,10}\b',
            # 驾驶证号
            r'\b[0-9]{12}[0-9Xx]\b',
            # 军官证号
            r'\b[军海空战陆海政][A-Z0-9]{6,12}\b',
            # 港澳通行证
            r'\b[W][0-9]{8,10}\b',
            # 台湾通行证
            r'\b[T][0-9]{8,10}\b',
            # IP地址
            r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b',
            # 数据库连接字符串
            r'(mongodb|mysql|postgres|redis):\/\/[^\s]+',
            # AWS密钥
            r'AKIA[0-9A-Z]{16}',
            # GitHub Token
            r'gh[pousr]_[A-Za-z0-9]{36,}',
            # JWT Token
            r'eyJ[A-Za-z0-9-_]+\.eyJ[A-Za-z0-9-_]+\.[A-Za-z0-9-_]+',
            # 比特币地址
            r'\b[13][a-km-zA-HJ-NP-Z1-9]{25,34}\b',
            # 支付宝/微信支付交易号
            r'\b20[0-9]{16,22}\b',
        ]
        
        text_lower = text.lower()
        text_stripped = text.strip()
        
        # 检查是否匹配敏感模式
        for pattern in sensitive_patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        
        # 检查是否是过短的内容（可能是密码）
        if 4 <= len(text_stripped) <= 16 and (text_stripped.isdigit() or text_stripped.isalpha()):
            return True
        
        # 检查是否只包含特殊字符（可能是密码）
        if 4 <= len(text_stripped) <= 20 and not text_stripped.isalnum():
            special_count = sum(1 for c in text_stripped if c in '!@#$%^&*()_+-=[]{}|;:,.<>?')
            if special_count >= 2:
                return True
        
        return False

    def _copy_clipboard_item(self, item):
        text = self.clipboard_history[self.clipboard_list.row(item)]
        self.clipboard.setText(text)
        self.show_toast("已复制到剪贴板")

    def _clear_clipboard_history(self):
        self.clipboard_history = []
        self.clipboard_list.clear()
        self.show_toast("已清空历史")

    # ================== 2. 批量文件处理 ==================
    def open_batch_processing(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("批量文件处理")
        dialog.setMinimumSize(700, 500)
        layout = QVBoxLayout(dialog)

        tabs = QTabWidget()
        rename_tab = QWidget()
        rename_layout = QVBoxLayout(rename_tab)
        self.rename_list = QListWidget()
        rename_layout.addWidget(self.rename_list)
        btn_row1 = QHBoxLayout()
        select_rename_btn = QPushButton(" 选择文件")
        select_rename_btn.setIcon(qta.icon('fa5s.folder-open', color='white'))
        select_rename_btn.setIconSize(ICON_SIZE_SMALL)
        select_rename_btn.clicked.connect(self._select_rename_files)
        btn_row1.addWidget(select_rename_btn)
        btn_row1.addStretch()
        rename_layout.addLayout(btn_row1)
        form_row = QHBoxLayout()
        form_row.addWidget(QLabel("前缀:"))
        self.prefix_edit = QLineEdit()
        form_row.addWidget(self.prefix_edit)
        form_row.addWidget(QLabel("起始序号:"))
        self.start_num = QSpinBox()
        self.start_num.setMinimum(1)
        form_row.addWidget(self.start_num)
        form_row.addWidget(QLabel("新后缀:"))
        self.new_ext_combo = QComboBox()
        self.new_ext_combo.addItems(["保持原后缀", ".txt", ".doc", ".docx", ".pdf", ".jpg", ".png", ".gif", ".bmp"])
        form_row.addWidget(self.new_ext_combo)
        rename_layout.addLayout(form_row)
        preview_btn = QPushButton(" 预览")
        preview_btn.setIcon(qta.icon('fa5s.eye', color='white'))
        preview_btn.clicked.connect(self._preview_rename)
        rename_layout.addWidget(preview_btn)
        exec_btn = QPushButton(" 执行重命名")
        exec_btn.setIcon(qta.icon('fa5s.check', color='white'))
        exec_btn.clicked.connect(self._exec_rename)
        rename_layout.addWidget(exec_btn)
        tabs.addTab(rename_tab, "批量重命名")

        suffix_tab = QWidget()
        suffix_layout = QVBoxLayout(suffix_tab)
        self.suffix_list = QListWidget()
        suffix_layout.addWidget(self.suffix_list)
        btn_row2 = QHBoxLayout()
        select_suffix_btn = QPushButton(" 选择文件")
        select_suffix_btn.setIcon(qta.icon('fa5s.folder-open', color='white'))
        select_suffix_btn.clicked.connect(self._select_suffix_files)
        btn_row2.addWidget(select_suffix_btn)
        btn_row2.addStretch()
        suffix_layout.addLayout(btn_row2)
        form_row2 = QHBoxLayout()
        form_row2.addWidget(QLabel("原后缀:"))
        self.old_ext_combo = QComboBox()
        self.old_ext_combo.addItems([".txt", ".doc", ".docx", ".pdf", ".jpg", ".png", ".gif", ".bmp", ".xlsx", ".csv"])
        form_row2.addWidget(self.old_ext_combo)
        form_row2.addWidget(QLabel("新后缀:"))
        self.new_ext_suffix = QComboBox()
        self.new_ext_suffix.addItems([".txt", ".doc", ".docx", ".pdf", ".jpg", ".png", ".gif", ".bmp", ".xlsx", ".csv"])
        form_row2.addWidget(self.new_ext_suffix)
        suffix_layout.addLayout(form_row2)
        suffix_exec_btn = QPushButton(" 批量修改后缀")
        suffix_exec_btn.setIcon(qta.icon('fa5s.sync-alt', color='white'))
        suffix_exec_btn.clicked.connect(self._batch_change_suffix)
        suffix_layout.addWidget(suffix_exec_btn)
        tabs.addTab(suffix_tab, "批量改后缀")
        
        replace_tab = QWidget()
        replace_layout = QVBoxLayout(replace_tab)
        
        replace_intro = QLabel("在多个文本文件中查找并替换指定字符串")
        replace_intro.setStyleSheet("color: #606266; padding: 5px 0;")
        replace_layout.addWidget(replace_intro)
        
        replace_file_layout = QHBoxLayout()
        replace_file_layout.addWidget(QLabel("选择文件:"))
        self.replace_file_list = QListWidget()
        self.replace_file_list.setMaximumHeight(80)
        replace_layout.addWidget(self.replace_file_list)
        
        select_replace_files_btn = QPushButton(" 选择文本文件")
        select_replace_files_btn.setIcon(qta.icon('fa5s.file-alt', color='#409eff'))
        select_replace_files_btn.clicked.connect(self._select_replace_files)
        replace_file_layout.addWidget(select_replace_files_btn)
        replace_layout.addLayout(replace_file_layout)
        
        find_layout = QHBoxLayout()
        find_layout.addWidget(QLabel("查找:"))
        self.replace_find_edit = QLineEdit()
        self.replace_find_edit.setPlaceholderText("输入要查找的字符串")
        find_layout.addWidget(self.replace_find_edit)
        replace_layout.addLayout(find_layout)
        
        repl_layout = QHBoxLayout()
        repl_layout.addWidget(QLabel("替换:"))
        self.replace_repl_edit = QLineEdit()
        self.replace_repl_edit.setPlaceholderText("输入替换后的字符串（留空则删除）")
        repl_layout.addWidget(self.replace_repl_edit)
        replace_layout.addLayout(repl_layout)
        
        self.replace_regex_check = QCheckBox("使用正则表达式")
        replace_layout.addWidget(self.replace_regex_check)
        
        btn_replace_row = QHBoxLayout()
        self.replace_preview_btn = QPushButton(" 预览")
        self.replace_preview_btn.setIcon(qta.icon('fa5s.eye', color='#409eff'))
        self.replace_preview_btn.clicked.connect(self._preview_replace)
        btn_replace_row.addWidget(self.replace_preview_btn)
        
        self.replace_exec_btn = QPushButton(" 执行替换")
        self.replace_exec_btn.setIcon(qta.icon('fa5s.check', color='white'))
        self.replace_exec_btn.setStyleSheet("background: #67c23a;")
        self.replace_exec_btn.clicked.connect(self._exec_replace)
        btn_replace_row.addWidget(self.replace_exec_btn)
        replace_layout.addLayout(btn_replace_row)
        
        self.replace_result = QTextBrowser()
        self.replace_result.setMaximumHeight(100)
        replace_layout.addWidget(QLabel("处理结果:"))
        replace_layout.addWidget(self.replace_result)
        
        tabs.addTab(replace_tab, "批量内容替换")

        layout.addWidget(tabs)
        dialog.exec()
    
    def _select_replace_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文本文件", "", "文本文件 (*.txt *.csv *.json *.xml *.html *.md *.log *.conf *.ini *.yaml *.yml *.py *.js *.java *.c *.cpp *.h *.cs *.go *.ts *.jsx *.tsx);;所有文件 (*.*)")
        self.replace_files = files
        self.replace_file_list.clear()
        for f in files:
            self.replace_file_list.addItem(f)
    
    def _preview_replace(self):
        if not hasattr(self, 'replace_files') or not self.replace_files:
            self.show_warning("提示", "请先选择文件")
            return
        find_text = self.replace_find_edit.text()
        if not find_text:
            self.show_warning("提示", "请输入要查找的字符串")
            return
        repl_text = self.replace_repl_edit.text()
        use_regex = self.replace_regex_check.isChecked()
        
        results = []
        total_matches = 0
        for f in self.replace_files:
            try:
                with open(f, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                if use_regex:
                    import re
                    matches = re.findall(find_text, content)
                else:
                    matches = content.count(find_text)
                total_matches += len(matches) if isinstance(matches, list) else matches
                results.append(f"📄 {os.path.basename(f)}: 找到 {len(matches) if isinstance(matches, list) else matches} 处")
            except Exception as e:
                logger.error(f"查找文件 {os.path.basename(f)} 失败: {e}", exc_info=True)
                results.append(f"❌ {os.path.basename(f)}: {str(e)[:30]}")
        
        self.replace_result.clear()
        self.replace_result.append(f"共找到 {total_matches} 处匹配\n")
        self.replace_result.append("\n".join(results[:20]))
        if len(results) > 20:
            self.replace_result.append(f"...还有 {len(results) - 20} 个文件")
    
    def _exec_replace(self):
        if not hasattr(self, 'replace_files') or not self.replace_files:
            self.show_warning("提示", "请先选择文件")
            return
        find_text = self.replace_find_edit.text()
        if not find_text:
            self.show_warning("提示", "请输入要查找的字符串")
            return
        
        repl_text = self.replace_repl_edit.text()
        use_regex = self.replace_regex_check.isChecked()
        
        success = 0
        skipped = 0
        failed = 0
        total_changes = 0
        
        for f in self.replace_files:
            try:
                with open(f, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                
                if use_regex:
                    import re
                    try:
                        new_content, count = re.subn(find_text, repl_text, content)
                    except re.error as e:
                        self.show_toast(f"正则表达式语法错误: {str(e)}")
                        return
                else:
                    count = content.count(find_text)
                    new_content = content.replace(find_text, repl_text)
                
                if count > 0:
                    with open(f, 'w', encoding='utf-8', errors='ignore') as file:
                        file.write(new_content)
                    success += 1
                    total_changes += count
                else:
                    skipped += 1
            except Exception as e:
                logger.error(f"替换文件 {os.path.basename(f)} 失败: {e}", exc_info=True)
                failed += 1
        
        result_msg = f"处理完成！\n✅ 成功: {success} 个文件，修改 {total_changes} 处\n"
        if skipped > 0:
            result_msg += f"⏭️ 跳过: {skipped} 个文件（无匹配）\n"
        if failed > 0:
            result_msg += f"❌ 失败: {failed} 个文件"
        self.replace_result.clear()
        self.replace_result.setText(result_msg)
        self.show_toast(result_msg)

    def _select_rename_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件")
        self.rename_files = files
        self.rename_list.clear()
        for f in files:
            self.rename_list.addItem(f)

    def _select_suffix_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件")
        self.suffix_files = files
        self.suffix_list.clear()
        for f in files:
            self.suffix_list.addItem(f)

    def _preview_rename(self):
        if not hasattr(self, 'rename_files'):
            return
        self.rename_list.clear()
        for i, f in enumerate(self.rename_files):
            ext = os.path.splitext(f)[1]
            new_ext = self.new_ext_combo.currentText()
            if new_ext != "保持原后缀":
                ext = new_ext
            new_name = f"{self.prefix_edit.text()}{self.start_num.value() + i}{ext}"
            self.rename_list.addItem(f"{os.path.basename(f)} → {new_name}")

    def _exec_rename(self):
        if not hasattr(self, 'rename_files'):
            return
        success = 0
        for i, f in enumerate(self.rename_files):
            try:
                dir_name = os.path.dirname(f)
                ext = os.path.splitext(f)[1]
                new_ext = self.new_ext_combo.currentText()
                if new_ext != "保持原后缀":
                    ext = new_ext
                new_name = f"{self.prefix_edit.text()}{self.start_num.value() + i}{ext}"
                new_path = os.path.join(dir_name, new_name)
                os.rename(f, new_path)
                success += 1
            except Exception as e:
                print(e)
        self.show_toast(f"成功重命名 {success} 个文件")

    def _batch_change_suffix(self):
        if not hasattr(self, 'suffix_files') or not self.suffix_files:
            self.show_warning("提示", "请先选择文件")
            return

        old = self.old_ext_combo.currentText()
        new = self.new_ext_suffix.currentText()

        if old == new:
            self.show_warning("提示", "新旧后缀相同，无需修改")
            return

        success = 0
        skip = 0
        failed = 0

        for f in self.suffix_files:
            try:
                file_ext = os.path.splitext(f)[1].lower()
                if file_ext != old.lower():
                    skip += 1
                    continue
                base_name = os.path.splitext(f)[0]
                new_path = base_name + new
                os.rename(f, new_path)
                success += 1
            except Exception as e:
                failed += 1
                print(f"修改失败 {os.path.basename(f)}: {str(e)}")

        self.suffix_list.clear()
        if success > 0:
            self.suffix_list.addItem(f"✅ 成功: {success} 个文件")
        if skip > 0:
            self.suffix_list.addItem(f"⏭️ 跳过: {skip} 个文件（扩展名不匹配）")
        if failed > 0:
            self.suffix_list.addItem(f"❌ 失败: {failed} 个文件")

        message = f"处理完成！\n成功: {success} 个\n"
        if skip > 0:
            message += f"跳过: {skip} 个\n"
        if failed > 0:
            message += f"失败: {failed} 个"
        self.show_toast(message)

    # ================== 3. 图片工具 ==================
    def open_image_tools(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("图片工具")
        dialog.setMinimumSize(600, 450)
        layout = QVBoxLayout(dialog)

        self.img_label = QLabel("点击选择图片")
        self.img_label.setAlignment(Qt.AlignCenter)
        self.img_label.setStyleSheet("border: 2px dashed #ccc; padding: 40px;")
        self.img_label.mousePressEvent = lambda e: self._select_ocr_image(self.img_label)
        layout.addWidget(self.img_label)

        opt_layout = QHBoxLayout()
        opt_layout.addWidget(QLabel("压缩质量:"))
        self.quality_spin = QSpinBox()
        self.quality_spin.setRange(1, 100)
        self.quality_spin.setValue(80)
        opt_layout.addWidget(self.quality_spin)
        
        opt_layout.addWidget(QLabel("格式:"))
        self.format_combo = QComboBox()
        self.format_combo.addItems(["JPG", "PNG", "GIF", "BMP", "WEBP", "TIFF", "ICO"])
        opt_layout.addWidget(self.format_combo)
        
        opt_layout.addWidget(QLabel("尺寸:"))
        self.size_combo = QComboBox()
        self.size_combo.addItems(["原图", "640x480", "800x600", "1024x768", "1280x720", "1920x1080"])
        opt_layout.addWidget(self.size_combo)
        
        # 宽度和高度输入（始终可用）
        self.custom_width_spin = QSpinBox()
        self.custom_width_spin.setRange(1, 8192)
        self.custom_width_spin.setValue(800)
        opt_layout.addWidget(QLabel("宽:"))
        opt_layout.addWidget(self.custom_width_spin)
        
        self.custom_height_spin = QSpinBox()
        self.custom_height_spin.setRange(1, 8192)
        self.custom_height_spin.setValue(600)
        opt_layout.addWidget(QLabel("高:"))
        opt_layout.addWidget(self.custom_height_spin)
        
        layout.addLayout(opt_layout)

        # 格式切换时更新尺寸选项
        def on_format_changed(text):
            self.size_combo.blockSignals(True)
            self.size_combo.clear()
            if text == "ICO":
                self.size_combo.addItems(["ICO 16x16", "ICO 32x32", "ICO 48x48", "ICO 64x64", "ICO 128x128", "ICO 256x256"])
                self.size_combo.setCurrentText("ICO 256x256")
            else:
                self.size_combo.addItems(["原图", "640x480", "800x600", "1024x768", "1280x720", "1920x1080"])
                self.size_combo.setCurrentText("原图")
            self.size_combo.blockSignals(False)
        
        # 尺寸选择联动
        def on_size_changed(text):
            is_ico_size = text.startswith("ICO ")
            if is_ico_size:
                self.format_combo.blockSignals(True)
                self.format_combo.setCurrentText("ICO")
                self.format_combo.blockSignals(False)
                ico_size_str = text.replace("ICO ", "")
                ico_w, ico_h = map(int, ico_size_str.split("x"))
                self.custom_width_spin.setValue(ico_w)
                self.custom_height_spin.setValue(ico_h)
        
        self.format_combo.currentTextChanged.connect(on_format_changed)
        self.size_combo.currentTextChanged.connect(on_size_changed)

        compress_btn = QPushButton(" 转换并保存")
        compress_btn.setIcon(qta.icon('fa5s.save', color='white'))
        compress_btn.clicked.connect(self._compress_image)
        layout.addWidget(compress_btn)
        dialog.exec()

    def _select_ocr_image(self, label):
        file, _ = QFileDialog.getOpenFileName(self, "选择图片", "", "图片文件 (*.png *.jpg *.jpeg *.bmp)")
        if file:
            self._ocr_image = file
            pixmap = QPixmap(file)
            label.setPixmap(pixmap.scaled(400, 300, Qt.KeepAspectRatio))
            label.setText("")

    def _compress_image(self):
        if not hasattr(self, "_ocr_image"):
            self.show_toast("请先选择图片")
            return
        img = QImage(self._ocr_image)
        if img.isNull():
            self.show_toast("图片加载失败")
            return
        fmt = self.format_combo.currentText()
        if fmt == "ICO":
            save_path, _ = QFileDialog.getSaveFileName(self, "保存ICO图标", "", "ICO文件 (*.ico)")
            if save_path:
                # 获取用户选择的ICO尺寸
                ico_size_str = self.size_combo.currentText().replace("ICO ", "")
                ico_w, ico_h = map(int, ico_size_str.split("x"))
                # ICO需要特殊处理，保持原始宽高比并填充透明背景
                ico_img = img.scaled(ico_w, ico_h, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                # 创建带透明背景的目标图像
                final_img = QImage(ico_w, ico_h, QImage.Format_ARGB32)
                final_img.fill(Qt.transparent)
                painter = QPainter(final_img)
                # 居中绘制缩放后的图像
                x = (ico_w - ico_img.width()) // 2
                y = (ico_h - ico_img.height()) // 2
                painter.drawImage(x, y, ico_img)
                painter.end()
                if final_img.save(save_path, "ICO"):
                    self.show_toast(f"✅ ICO ({ico_size_str})已保存")
            return
        
        # 处理非ICO格式的尺寸调整
        size_str = self.size_combo.currentText()
        if size_str != "原图" and not size_str.startswith("ICO "):
            w, h = map(int, size_str.split("x"))
            img = img.scaled(w, h, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        else:
            # 使用用户输入的宽高
            w = self.custom_width_spin.value()
            h = self.custom_height_spin.value()
            if size_str == "原图":
                img = img.scaled(w, h, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        
        save_path, _ = QFileDialog.getSaveFileName(self, "保存图片", "", f"{fmt}文件 (*.{fmt.lower()})")
        if save_path:
            if img.save(save_path, fmt.upper(), self.quality_spin.value()):
                self.show_toast(f"✅ {fmt}已保存")

    # ================== 4. 压缩解压 ==================
    def open_zip_tools(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("压缩解压工具")
        dialog.setMinimumSize(550, 450)
        layout = QVBoxLayout(dialog)
        tabs = QTabWidget()
        compress_tab = QWidget()
        compress_layout = QVBoxLayout(compress_tab)

        compress_label = QLabel("拖放文件到下方，或点击添加文件按钮：")
        compress_layout.addWidget(compress_label)

        self.compress_list = DragDropListWidget()
        self.compress_list.setMinimumHeight(180)
        self.compress_list.setStyleSheet("""
            QListWidget { 
                border: 2px dashed #ccc; 
                border-radius: 8px; 
                padding: 10px; 
                background: #fafafa;
            }
            QListWidget::item { 
                padding: 5px; 
            }
        """)
        compress_layout.addWidget(self.compress_list)

        compress_btn_row = QHBoxLayout()
        add_btn = QPushButton(" 添加文件")
        add_btn.setIcon(qta.icon('fa5s.plus', color='white'))
        add_btn.clicked.connect(self._add_zip_files)
        compress_btn_row.addWidget(add_btn)
        
        del_btn = QPushButton(" 删除选中")
        del_btn.setIcon(qta.icon('fa5s.trash-alt', color='white'))
        del_btn.clicked.connect(self._delete_selected_zip_files)
        compress_btn_row.addWidget(del_btn)
        compress_btn_row.addStretch()
        compress_layout.addLayout(compress_btn_row)

        create_btn = QPushButton(" 创建压缩包")
        create_btn.setIcon(qta.icon('fa5s.file-archive', color='white'))
        create_btn.clicked.connect(self._create_zip)
        compress_layout.addWidget(create_btn)
        tabs.addTab(compress_tab, "压缩")

        extract_tab = QWidget()
        extract_layout = QVBoxLayout(extract_tab)

        extract_label = QLabel("拖放压缩包到下方，或点击选择压缩包按钮：")
        extract_layout.addWidget(extract_label)

        self.extract_list = DragDropListWidget()
        self.extract_list.setMinimumHeight(180)
        self.extract_list.setStyleSheet("""
            QListWidget { 
                border: 2px dashed #ccc; 
                border-radius: 8px; 
                padding: 10px; 
                background: #fafafa;
            }
            QListWidget::item { 
                padding: 5px; 
            }
        """)
        extract_layout.addWidget(self.extract_list)

        select_zip_btn = QPushButton(" 选择压缩包")
        select_zip_btn.setIcon(qta.icon('fa5s.folder-open', color='white'))
        select_zip_btn.clicked.connect(self._select_zip_files)
        extract_layout.addWidget(select_zip_btn)

        extract_btn = QPushButton(" 解压到...")
        extract_btn.setIcon(qta.icon('fa5s.folder-open', color='white'))
        extract_btn.clicked.connect(self._extract_zips)
        extract_layout.addWidget(extract_btn)
        tabs.addTab(extract_tab, "解压")
        layout.addWidget(tabs)
        dialog.exec()

    def _add_zip_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件")
        for f in files:
            if f not in self.compress_list.file_list:
                self.compress_list.file_list.append(f)
                self.compress_list.addItem(f)

    def _delete_selected_zip_files(self):
        selected_items = self.compress_list.selectedItems()
        if not selected_items:
            self.show_toast("请先选中要删除的文件")
            return
        for item in selected_items:
            file_path = item.text()
            if file_path in self.compress_list.file_list:
                self.compress_list.file_list.remove(file_path)
            row = self.compress_list.row(item)
            self.compress_list.takeItem(row)

    def _create_zip(self):
        zip_files = self.compress_list.get_files()
        if not zip_files:
            self.show_toast("请先添加文件")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存压缩包", "", "Zip文件 (*.zip)")
        if save_path:
            with zipfile.ZipFile(save_path, 'w') as zf:
                for f in zip_files:
                    zf.write(f, os.path.basename(f))
            self.show_toast("压缩包已创建")

    def _select_zip_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择压缩包", "", "Zip文件 (*.zip)")
        for f in files:
            if f not in self.extract_list.file_list:
                self.extract_list.file_list.append(f)
                self.extract_list.addItem(f)

    def _extract_zips(self):
        zip_files = self.extract_list.get_files()
        if not zip_files:
            self.show_toast("请先选择压缩包")
            return
        folder = QFileDialog.getExistingDirectory(self, "选择解压目录")
        if folder:
            for zip_path in zip_files:
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    zf.extractall(folder)
            self.show_toast("解压完成")

    # ================== 5. 批量文字提取 ==================
    def open_batch_text_extract(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("批量文字提取")
        dialog.setMinimumSize(800, 600)
        layout = QVBoxLayout(dialog)

        file_list_label = QLabel("已选文件：")
        layout.addWidget(file_list_label)
        self.extract_file_list = QTableWidget()
        self.extract_file_list.setColumnCount(3)
        self.extract_file_list.setHorizontalHeaderLabels(['文件名', '状态', '错误信息'])
        self.extract_file_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.extract_file_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.extract_file_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.extract_file_list.setAlternatingRowColors(True)
        self.extract_file_list.setSelectionBehavior(QTableWidget.SelectRows)
        layout.addWidget(self.extract_file_list)

        btn_row = QHBoxLayout()
        select_files_btn = QPushButton(" 选择文件")
        select_files_btn.setIcon(qta.icon('fa5s.folder-open', color='white'))
        select_files_btn.setIconSize(ICON_SIZE_SMALL)
        select_files_btn.clicked.connect(self._select_extract_files)
        btn_row.addWidget(select_files_btn)

        clear_btn = QPushButton(" 清空列表")
        clear_btn.setIcon(qta.icon('fa5s.trash-alt', color='white'))
        clear_btn.clicked.connect(lambda: self.extract_file_list.setRowCount(0))
        btn_row.addWidget(clear_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        output_label = QLabel("输出设置：")
        layout.addWidget(output_label)
        output_row = QHBoxLayout()
        output_row.addWidget(QLabel("保存位置:"))
        self.extract_output_path = QLineEdit()
        self.extract_output_path.setPlaceholderText("选择一个文件夹保存提取的文字...")
        output_row.addWidget(self.extract_output_path)
        select_output_btn = QPushButton(" 浏览")
        select_output_btn.clicked.connect(self._select_extract_output)
        output_row.addWidget(select_output_btn)
        layout.addLayout(output_row)

        self.merge_output_check = QCheckBox("合并所有文件到一个文本文件")
        self.merge_output_check.setChecked(True)
        layout.addWidget(self.merge_output_check)

        extract_btn = QPushButton(" 开始提取文字")
        extract_btn.setIcon(qta.icon('fa5s.file-alt', color='white'))
        extract_btn.clicked.connect(self._do_batch_extract)
        layout.addWidget(extract_btn)

        hint_label = QLabel("支持: Word (.docx) | Excel (.xlsx) | PDF | TXT | 图片 (OCR)")
        hint_label.setStyleSheet("color: #6b7280; font-size: 11px; padding: 8px; background-color: #f3f4f6; border-radius: 4px;")
        layout.addWidget(hint_label)

        dialog.exec()

    def _select_extract_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件", "",
            "所有支持文件 (*.docx *.xlsx *.pdf *.txt *.png *.jpg *.jpeg *.bmp);;"
            "Word文件 (*.docx);;Excel文件 (*.xlsx);;PDF文件 (*.pdf);;"
            "文本文件 (*.txt);;图片文件 (*.png *.jpg *.jpeg *.bmp)")
        for f in files:
            row = self.extract_file_list.rowCount()
            self.extract_file_list.insertRow(row)
            self.extract_file_list.setItem(row, 0, QTableWidgetItem(os.path.basename(f)))
            self.extract_file_list.setItem(row, 1, QTableWidgetItem("待处理"))
            self.extract_file_list.setItem(row, 2, QTableWidgetItem(""))
            self.extract_file_list.item(row, 0).setData(Qt.UserRole, f)

    def _select_extract_output(self):
        folder = QFileDialog.getExistingDirectory(self, "选择保存位置")
        if folder:
            self.extract_output_path.setText(folder)

    def _do_batch_extract(self):
        if self.extract_file_list.rowCount() == 0:
            self.show_toast("请先选择要提取的文件")
            return
        output_folder = self.extract_output_path.text()
        if not output_folder:
            self.show_toast("请选择保存位置")
            return

        os.makedirs(output_folder, exist_ok=True)
        success_count = 0
        error_count = 0

        merged_text = ""

        for i in range(self.extract_file_list.rowCount()):
            file_path = self.extract_file_list.item(i, 0).data(Qt.UserRole)
            file_name = os.path.basename(file_path)
            self.extract_file_list.setItem(i, 1, QTableWidgetItem("处理中..."))
            self.extract_file_list.item(i, 1).setBackground(QColor("#fff3cd"))
            QApplication.processEvents()
            
            try:
                text = self._extract_text_from_file(file_path)
                if text and not (text.startswith("[失败") or text.startswith("[Word提取失败") or text.startswith("[Excel提取失败")):
                    file_header = f"【{file_name}】\n" + "=" * 50 + "\n\n"
                    if self.merge_output_check.isChecked():
                        merged_text += file_header + text + "\n\n"
                    else:
                        base_name = os.path.splitext(file_name)[0]
                        output_file = os.path.join(output_folder, f"{base_name}_提取文字.txt")
                        with open(output_file, 'w', encoding='utf-8') as f:
                            f.write(file_header + text)
                    success_count += 1
                    self.extract_file_list.setItem(i, 1, QTableWidgetItem("✓ 成功"))
                    self.extract_file_list.item(i, 1).setBackground(QColor("#d4edda"))
                    self.extract_file_list.setItem(i, 2, QTableWidgetItem(""))
                else:
                    error_count += 1
                    error_msg = text if text else "提取失败"
                    self.extract_file_list.setItem(i, 1, QTableWidgetItem("✗ 失败"))
                    self.extract_file_list.item(i, 1).setBackground(QColor("#f8d7da"))
                    self.extract_file_list.setItem(i, 2, QTableWidgetItem(error_msg))
            except Exception as e:
                error_count += 1
                error_msg = str(e)
                self.extract_file_list.setItem(i, 1, QTableWidgetItem("✗ 失败"))
                self.extract_file_list.item(i, 1).setBackground(QColor("#f8d7da"))
                self.extract_file_list.setItem(i, 2, QTableWidgetItem(error_msg))

        if self.merge_output_check.isChecked() and merged_text:
            output_file = os.path.join(output_folder, "合并提取文本.txt")
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("批量文字提取结果\n")
                f.write("=" * 60 + "\n\n")
                f.write(merged_text)

        self.show_toast(f"提取完成！成功: {success_count}, 失败: {error_count}")

    def _extract_text_from_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext in ['.docx']:
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for p in doc.paragraphs:
                    text += p.text + "\n"
                for table in doc.tables:
                    text += "\n[表格]\n"
                    for row in table.rows:
                        row_text = "\t".join([cell.text for cell in row.cells])
                        text += row_text + "\n"
                return text.strip()
            except Exception as e:
                return f"[Word提取失败: {str(e)}]"
        elif ext in ['.xlsx']:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                text = ""
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    text += f"\n【{sheet}】\n"
                    for row in ws.iter_rows(values_only=True):
                        text += '\t'.join([str(cell) if cell else '' for cell in row]) + '\n'
                return text
            except:
                return "[Excel提取失败]"
        elif ext in ['.png', '.jpg', '.jpeg', '.bmp']:
            try:
                import pytesseract
                from PIL import Image
                import subprocess
                import tempfile

                base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                builtin_tesseract = os.path.join(base_dir, "tesseract", "tesseract.exe")
                builtin_tessdata = os.path.join(base_dir, "tesseract", "tessdata")

                tesseract_cmd = None
                tessdata_dir = None
                if os.path.exists(builtin_tesseract):
                    tesseract_cmd = builtin_tesseract
                    if os.path.exists(builtin_tessdata):
                        tessdata_dir = builtin_tessdata
                else:
                    tesseract_paths = [
                        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                        r'C:\Tesseract-OCR\tesseract.exe',
                        r'C:\Users\%USERNAME%\AppData\Local\Tesseract-OCR\tesseract.exe'
                    ]
                    for path in tesseract_paths:
                        expanded_path = os.path.expandvars(path)
                        if os.path.exists(expanded_path):
                            tesseract_cmd = expanded_path
                            break

                if not tesseract_cmd:
                    return "[未找到Tesseract OCR引擎]"

                img = Image.open(file_path)
                img = img.convert('L')
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    img.save(tmp.name)
                    tmp_path = tmp.name
                try:
                    env = os.environ.copy()
                    if tessdata_dir:
                        env['TESSDATA_PREFIX'] = tessdata_dir
                    result = subprocess.run(
                        [tesseract_cmd, tmp_path, 'stdout', '-l', 'chi_sim+eng', '--psm', '6'],
                        env=env,
                        capture_output=True,
                        text=False
                    )
                    text = ""
                    if result.stdout:
                        try:
                            text = result.stdout.decode('utf-8', errors='ignore')
                        except:
                            try:
                                text = result.stdout.decode('gbk', errors='ignore')
                            except:
                                try:
                                    text = result.stdout.decode('gb18030', errors='ignore')
                                except:
                                    text = ""
                    return text if text.strip() else "[未识别到文字]"
                finally:
                    os.unlink(tmp_path)
            except Exception as e:
                return f"[图片OCR失败: {str(e)}]"
        else:
            return "[不支持的文件格式]"

    # ================== 6. 图片OCR ==================
    def open_image_ocr(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("图片文字识别")
        dialog.setMinimumSize(600, 500)
        layout = QVBoxLayout(dialog)
        
        theme = {}
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            theme = self.main_window.theme_manager.get_theme()
        
        self.ocr_img_label = QLabel("点击选择图片")
        self.ocr_img_label.setAlignment(Qt.AlignCenter)
        self.ocr_img_label.setStyleSheet("border: 2px dashed #ccc; padding: 40px;")
        self.ocr_img_label.mousePressEvent = lambda e: self._select_ocr_image(self.ocr_img_label)
        layout.addWidget(self.ocr_img_label)
        self.ocr_result = QTextEdit()
        self.ocr_result.setReadOnly(True)
        layout.addWidget(self.ocr_result)
        btn_row = QHBoxLayout()
        ocr_btn = QPushButton(" 识别文字")
        ocr_btn.setIcon(qta.icon('fa5s.magic', color='white'))
        ocr_btn.clicked.connect(lambda: self._start_ocr_with_progress(dialog))
        copy_btn = QPushButton(" 复制结果")
        copy_btn.setIcon(qta.icon('fa5s.copy', color='white'))
        copy_btn.clicked.connect(self._copy_ocr_result)
        self.ocr_cancel_btn = QPushButton(" 取消识别")
        self.ocr_cancel_btn.setIcon(qta.icon('fa5s.times', color='white'))
        self.ocr_cancel_btn.clicked.connect(self._cancel_ocr)
        self.ocr_cancel_btn.setEnabled(False)
        btn_row.addWidget(ocr_btn)
        btn_row.addWidget(copy_btn)
        btn_row.addWidget(self.ocr_cancel_btn)
        layout.addLayout(btn_row)
        dialog.exec()

    def _start_ocr_with_progress(self, dialog):
        if not hasattr(self, "_ocr_image") or not self._ocr_image:
            self.show_toast("请先选择图片")
            return
        tesseract_cmd = self._get_tesseract_path()
        if not tesseract_cmd or not os.path.exists(tesseract_cmd):
            self.show_warning("错误", "未找到 Tesseract OCR 引擎")
            return
        tessdata_dir = self._get_tessdata_dir()
        
        # 清除之前的结果并显示准备状态
        self.ocr_result.clear()
        self.ocr_result.setText("准备开始识别...")
        self.ocr_cancel_btn.setEnabled(True)
        
        # 创建并启动OCR线程
        self._ocr_thread = OcrThread(self._ocr_image, tesseract_cmd, tessdata_dir)
        self._ocr_thread.progress.connect(self.update_ocr_progress)
        self._ocr_thread.log.connect(self.update_ocr_log)
        self._ocr_thread.finished.connect(lambda success, msg: self._on_ocr_finished(success, msg, dialog))
        self._ocr_thread.start()

    def _on_ocr_finished(self, success, message, ocr_dialog):
        self.ocr_cancel_btn.setEnabled(False)
        if success:
            self.ocr_result.setText(f"【识别完成】\n\n{message}")
        else:
            self.ocr_result.setText(f"【识别失败】\n\n{message}")
        self._ocr_thread = None

    def _cancel_ocr(self):
        if hasattr(self, '_ocr_thread') and self._ocr_thread:
            # 使用标志位优雅停止线程，而不是强制终止
            self._ocr_thread.cancel()
            self._ocr_thread.wait(5000)  # 最多等待5秒
            self._ocr_thread = None
        self.ocr_cancel_btn.setEnabled(False)
        self.ocr_result.append("\n【已取消识别】")

    def update_ocr_progress(self, value):
        """更新OCR进度到文本框"""
        current_text = self.ocr_result.toPlainText()
        # 如果最后一行是进度信息，则替换它
        lines = current_text.split('\n')
        if lines and '进度' in lines[-1]:
            lines[-1] = f"进度: {value}%"
            self.ocr_result.setText('\n'.join(lines))
        else:
            self.ocr_result.append(f"进度: {value}%")
        # 滚动到底部
        scrollbar = self.ocr_result.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    def update_ocr_log(self, msg):
        """更新OCR日志到文本框"""
        self.ocr_result.append(msg)
        # 滚动到底部
        scrollbar = self.ocr_result.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    def _get_tesseract_path(self):
        import sys
        import shutil
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        builtin_path = os.path.join(base_dir, "tesseract", "tesseract.exe")
        if os.path.exists(builtin_path):
            return builtin_path
        if getattr(sys, 'frozen', False):
            frozen_path = os.path.join(sys._MEIPASS, "tesseract", "tesseract.exe")
            if os.path.exists(frozen_path):
                return frozen_path
        system_path = shutil.which('tesseract')
        if system_path:
            return system_path
        custom_path, _ = QFileDialog.getOpenFileName(self, "请选择 Tesseract 可执行文件", "", "Tesseract 程序 (tesseract.exe tesseract)")
        return custom_path if custom_path else None

    def _get_tessdata_dir(self):
        import sys
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        builtin_path = os.path.join(base_dir, "tesseract", "tessdata")
        if os.path.exists(builtin_path):
            return builtin_path
        if getattr(sys, 'frozen', False):
            frozen_path = os.path.join(sys._MEIPASS, "tesseract", "tessdata")
            if os.path.exists(frozen_path):
                return frozen_path
        return None

    def _do_ocr(self):
        if not hasattr(self, "_ocr_image") or not self._ocr_image:
            self.show_toast("请先选择图片")
            return
        tesseract_cmd = self._get_tesseract_path()
        if not tesseract_cmd or not os.path.exists(tesseract_cmd):
            self.show_warning("错误", "未找到 Tesseract OCR 引擎")
            return
        try:
            import pytesseract
            from PIL import Image
            import subprocess
            import tempfile
            from PIL import ImageEnhance, ImageFilter

            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

            img = Image.open(self._ocr_image)
            if img.mode != 'L':
                img = img.convert('L')
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.5)

            psm_options = ['3', '6', '4', '1']
            recognized_text = ""
            best_text = ""
            tessdata_dir = self._get_tessdata_dir()

            for psm in psm_options:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    img.save(tmp_img.name, 'PNG')
                    tmp_img_path = tmp_img.name
                with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp_out:
                    tmp_out_path = tmp_out.name[:-4]
                try:
                    cmd = [tesseract_cmd]
                    if tessdata_dir and os.path.exists(tessdata_dir):
                        cmd.extend(['--tessdata-dir', tessdata_dir])
                    cmd.extend([tmp_img_path, tmp_out_path, '-l', 'chi_sim+eng', '--psm', psm])
                    result = subprocess.run(cmd, capture_output=True, text=False)
                    output_txt_path = tmp_out_path + '.txt'
                    if os.path.exists(output_txt_path):
                        try:
                            with open(output_txt_path, 'r', encoding='utf-8') as f:
                                recognized_text = f.read()
                        except UnicodeDecodeError:
                            try:
                                with open(output_txt_path, 'r', encoding='gbk') as f:
                                    recognized_text = f.read()
                            except:
                                with open(output_txt_path, 'r', encoding='gb18030', errors='ignore') as f:
                                    recognized_text = f.read()
                        if len(recognized_text.strip()) > len(best_text.strip()):
                            best_text = recognized_text
                finally:
                    if os.path.exists(tmp_img_path):
                        try:
                            os.unlink(tmp_img_path)
                        except:
                            pass
                    if os.path.exists(tmp_out_path + '.txt'):
                        try:
                            os.unlink(tmp_out_path + '.txt')
                        except:
                            pass
            final_text = best_text if best_text else recognized_text
            self.ocr_result.setText(final_text.strip() if final_text.strip() else "未识别到文字")
            self.show_toast("识别完成")
        except ImportError as e:
            self.ocr_result.setText(f"缺少依赖库: {str(e)}")
        except Exception as e:
            self.ocr_result.setText(f"识别失败: {str(e)}")
            traceback.print_exc()

    def _copy_ocr_result(self):
        text = self.ocr_result.toPlainText()
        if text:
            QApplication.clipboard().setText(text)
            self.show_toast("已复制到剪贴板")

    # ================== 7. 文本格式化 ==================
    def open_text_format(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("文本格式化")
        dialog.setMinimumSize(750, 600)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        input_label = QLabel("输入文本:")
        input_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        layout.addWidget(input_label)

        input_toolbar = QHBoxLayout()
        import_btn = QPushButton(" 导入文件")
        import_btn.setIcon(qta.icon('fa5s.file-import', color='#409eff'))
        import_btn.clicked.connect(self._import_text_file)
        clear_btn = QPushButton(" 清空")
        clear_btn.setIcon(qta.icon('fa5s.eraser', color='#f56c6c'))
        clear_btn.clicked.connect(lambda: self.input_text.clear())
        input_toolbar.addWidget(import_btn)
        input_toolbar.addWidget(clear_btn)
        input_toolbar.addStretch()
        layout.addLayout(input_toolbar)

        self.input_text = QTextEdit()
        self.input_text.setStyleSheet("border: 1px solid #dcdfe6; border-radius: 4px; padding: 8px;")
        self.input_text.setMinimumHeight(120)
        layout.addWidget(self.input_text)

        input_tip = QLabel("⚠️ 提示：此功能仅支持简单的文本编辑，复杂格式（如表格、图片、特殊样式等）导入时可能会丢失，建议先备份原文件。")
        input_tip.setStyleSheet("color: #e6a23c; font-size: 12px; background: #fdf6ec; padding: 6px 10px; border-radius: 4px; border: 1px solid #f5dab1;")
        layout.addWidget(input_tip)

        format_row = QHBoxLayout()
        format_row.addWidget(QLabel("格式化类型:"))
        self.format_type = QComboBox()
        self.format_type.addItems([
            "JSON", "SQL", "Markdown", "HTML", "XML", "普通文本",
            "文本去重", "文本排序", "文本反转", "去除多余空格", "去除空白行",
            "全角转半角", "半角转全角", "大小写转换",
            "添加行号", "添加自定义前缀", "添加换行符",
            "替换特定字符", "修复URL格式"
        ])
        format_row.addWidget(self.format_type, 1)

        format_btn = QPushButton(" 格式化")
        format_btn.setIcon(qta.icon('fa5s.code', color='white'))
        format_btn.setStyleSheet("background: #409eff; color: white; padding: 8px 20px; border-radius: 4px;")
        format_btn.clicked.connect(self._format_text)
        format_row.addWidget(format_btn)
        layout.addLayout(format_row)

        output_label = QLabel("结果:")
        output_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        layout.addWidget(output_label)

        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        self.output_text.setStyleSheet("border: 1px solid #dcdfe6; border-radius: 4px; padding: 8px; background: #f5f7fa;")
        self.output_text.setMinimumHeight(120)
        layout.addWidget(self.output_text)

        output_toolbar = QHBoxLayout()
        copy_btn = QPushButton(" 复制结果")
        copy_btn.setIcon(qta.icon('fa5s.copy', color='white'))
        copy_btn.clicked.connect(self._copy_formatted)
        save_btn = QPushButton(" 保存到文件")
        save_btn.setIcon(qta.icon('fa5s.save', color='white'))
        save_btn.clicked.connect(self._save_formatted_text)
        output_toolbar.addWidget(copy_btn)
        output_toolbar.addWidget(save_btn)
        output_toolbar.addStretch()
        layout.addLayout(output_toolbar)

        dialog.exec()

    def _import_text_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择文本文件",
            "",
            "文本文件 (*.txt *.json *.sql *.md *.html *.xml);;所有文件 (*.*)"
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.input_text.setText(content)
                self.show_toast("✅ 文件导入成功")
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        content = f.read()
                    self.input_text.setText(content)
                    self.show_toast("✅ 文件导入成功")
                except Exception as e:
                    self.show_error("错误", f"文件读取失败: {str(e)}")
            except Exception as e:
                self.show_error("错误", f"文件读取失败: {str(e)}")

    def _save_formatted_text(self):
        text = self.output_text.toPlainText()
        if not text:
            self.show_toast("没有可保存的内容")
            return
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存文件",
            "",
            "文本文件 (*.txt);;JSON文件 (*.json);;SQL文件 (*.sql);;Markdown (*.md);;所有文件 (*.*)"
        )
        if save_path:
            try:
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.show_toast("✅ 文件保存成功")
            except Exception as e:
                self.show_error("错误", f"文件保存失败: {str(e)}")

    def _format_text(self):
        text = self.input_text.toPlainText()
        if not text.strip():
            self.show_toast("请输入文本")
            return

        fmt = self.format_type.currentText()

        try:
            formatted = ""
            if fmt == "JSON":
                obj = json.loads(text)
                formatted = json.dumps(obj, indent=2, ensure_ascii=False)
            elif fmt == "SQL":
                sql = text.strip()
                keywords = ['SELECT', 'FROM', 'WHERE', 'AND', 'OR', 'ORDER BY', 'GROUP BY', 'HAVING', 'JOIN', 'LEFT', 'RIGHT', 'INNER', 'OUTER', 'INSERT', 'UPDATE', 'DELETE', 'SET', 'VALUES', 'CREATE', 'ALTER', 'DROP', 'TABLE', 'DATABASE']
                formatted = sql
                for kw in keywords:
                    formatted = formatted.replace(kw, '\n' + kw)
                formatted = '\n'.join([l.strip() for l in formatted.split('\n') if l.strip()])
            elif fmt == "Markdown":
                lines = [l.strip() for l in text.split('\n')]
                formatted = '\n\n'.join(lines)
            elif fmt == "HTML":
                formatted = text.replace('>', '>\n').replace('<', '\n<').strip()
                formatted = '\n'.join([l.strip() for l in formatted.split('\n') if l.strip()])
            elif fmt == "XML":
                formatted = text.replace('>', '>\n').replace('<', '\n<').strip()
                formatted = '\n'.join([l.strip() for l in formatted.split('\n') if l.strip()])
            elif fmt == "文本去重":
                lines = text.split('\n')
                seen = set()
                new_lines = []
                for line in lines:
                    if line not in seen:
                        seen.add(line)
                        new_lines.append(line)
                formatted = '\n'.join(new_lines)
            elif fmt == "文本排序":
                lines = text.split('\n')
                sorted_lines = sorted(lines, key=lambda x: x.strip().lower())
                formatted = '\n'.join(sorted_lines)
            elif fmt == "文本反转":
                formatted = text[::-1]
            elif fmt == "全角转半角":
                formatted = self._full_to_half(text)
            elif fmt == "半角转全角":
                formatted = self._half_to_full(text)
            elif fmt == "去除空白行":
                # 去除空白行
                lines = text.split('\n')
                formatted = '\n'.join(line for line in lines if line.strip())
            elif fmt == "添加行号":
                # 添加行号
                lines = text.split('\n')
                format_dialog = QDialog(self)
                format_dialog.setWindowTitle("设置行号格式")
                format_layout = QVBoxLayout(format_dialog)
                format_layout.addWidget(QLabel("选择行号格式:"))
                format_combo = QComboBox()
                format_combo.addItems(["1. 文本", "[1] 文本", "1) 文本"])
                format_combo.setCurrentIndex(0)
                format_layout.addWidget(format_combo)
                btn_box = QHBoxLayout()
                ok_btn = QPushButton("确定")
                cancel_btn = QPushButton("取消")
                btn_box.addStretch()
                btn_box.addWidget(ok_btn)
                btn_box.addWidget(cancel_btn)
                format_layout.addLayout(btn_box)
                
                user_confirmed = [False]
                selected_format = [0]
                
                def on_ok():
                    selected_format[0] = format_combo.currentIndex()
                    user_confirmed[0] = True
                    format_dialog.accept()
                
                ok_btn.clicked.connect(on_ok)
                cancel_btn.clicked.connect(format_dialog.reject)
                
                if format_dialog.exec() and user_confirmed[0]:
                    new_lines = []
                    fmt_type = selected_format[0]
                    for i, line in enumerate(lines, 1):
                        if fmt_type == 0:
                            new_lines.append(f"{i}. {line}")
                        elif fmt_type == 1:
                            new_lines.append(f"[{i}] {line}")
                        else:
                            new_lines.append(f"{i}) {line}")
                    formatted = '\n'.join(new_lines)
                else:
                    return
            elif fmt == "添加自定义前缀":
                # 添加自定义前缀
                prefix, ok = QInputDialog.getText(self, "设置前缀", "请输入前缀:")
                if ok and prefix:
                    lines = text.split('\n')
                    formatted = '\n'.join(f"{prefix}{line}" for line in lines)
                else:
                    return
            elif fmt == "去除多余空格":
                # 去除多余空格
                import re
                lines = text.split('\n')
                processed_lines = []
                for line in lines:
                    line = line.strip()
                    line = re.sub(r'\s+', ' ', line)
                    processed_lines.append(line)
                formatted = '\n'.join(processed_lines)
            elif fmt == "大小写转换":
                # 大小写转换
                case_dialog = QDialog(self)
                case_dialog.setWindowTitle("设置大小写格式")
                case_layout = QVBoxLayout(case_dialog)
                case_layout.addWidget(QLabel("选择转换方式:"))
                case_combo = QComboBox()
                case_combo.addItems(["全部大写", "全部小写", "首字母大写", "每个单词首字母大写"])
                case_combo.setCurrentIndex(0)
                case_layout.addWidget(case_combo)
                btn_box2 = QHBoxLayout()
                ok_btn2 = QPushButton("确定")
                cancel_btn2 = QPushButton("取消")
                btn_box2.addStretch()
                btn_box2.addWidget(ok_btn2)
                btn_box2.addWidget(cancel_btn2)
                case_layout.addLayout(btn_box2)
                
                user_confirmed2 = [False]
                selected_case = [0]
                
                def on_ok2():
                    selected_case[0] = case_combo.currentIndex()
                    user_confirmed2[0] = True
                    case_dialog.accept()
                
                ok_btn2.clicked.connect(on_ok2)
                cancel_btn2.clicked.connect(case_dialog.reject)
                
                if case_dialog.exec() and user_confirmed2[0]:
                    case_type = selected_case[0]
                    if case_type == 0:
                        formatted = text.upper()
                    elif case_type == 1:
                        formatted = text.lower()
                    elif case_type == 2:
                        formatted = text.capitalize()
                    else:
                        import re
                        formatted = re.sub(r'\b\w', lambda m: m.group(0).upper(), text.lower())
                else:
                    return
            elif fmt == "添加换行符":
                # 添加换行符
                new_line_dialog = QDialog(self)
                new_line_dialog.setWindowTitle("设置换行")
                new_line_layout = QVBoxLayout(new_line_dialog)
                new_line_layout.addWidget(QLabel("选择换行方式:"))
                new_line_combo = QComboBox()
                new_line_combo.addItems(["每行末尾添加换行", "每N个字符后换行", "在特定字符后换行"])
                new_line_combo.setCurrentIndex(0)
                new_line_layout.addWidget(new_line_combo)
                
                n_label = QLabel("N值(仅适用于第2种):")
                n_spin = QSpinBox()
                n_spin.setRange(1, 1000)
                n_spin.setValue(80)
                n_row = QHBoxLayout()
                n_row.addWidget(n_label)
                n_row.addWidget(n_spin)
                new_line_layout.addLayout(n_row)
                
                sep_label = QLabel("分隔符(仅适用于第3种):")
                sep_edit = QLineEdit(",")
                sep_row = QHBoxLayout()
                sep_row.addWidget(sep_label)
                sep_row.addWidget(sep_edit)
                new_line_layout.addLayout(sep_row)
                
                btn_box3 = QHBoxLayout()
                ok_btn3 = QPushButton("确定")
                cancel_btn3 = QPushButton("取消")
                btn_box3.addStretch()
                btn_box3.addWidget(ok_btn3)
                btn_box3.addWidget(cancel_btn3)
                new_line_layout.addLayout(btn_box3)
                
                user_confirmed3 = [False]
                selected_option = [0]
                n_value = [80]
                sep_value = [","]
                
                def on_ok3():
                    selected_option[0] = new_line_combo.currentIndex()
                    n_value[0] = n_spin.value()
                    sep_value[0] = sep_edit.text()
                    user_confirmed3[0] = True
                    new_line_dialog.accept()
                
                ok_btn3.clicked.connect(on_ok3)
                cancel_btn3.clicked.connect(new_line_dialog.reject)
                
                if new_line_dialog.exec() and user_confirmed3[0]:
                    opt = selected_option[0]
                    if opt == 0:
                        # 每行末尾添加换行
                        lines = text.split('\n')
                        formatted = '\n'.join(f"{line}\n" for line in lines).rstrip('\n')
                    elif opt == 1:
                        # 每N个字符后换行
                        n = n_value[0]
                        result = []
                        for i in range(0, len(text), n):
                            result.append(text[i:i+n])
                        formatted = '\n'.join(result)
                    else:
                        # 在特定字符后换行
                        sep = sep_value[0]
                        formatted = text.replace(sep, sep + '\n')
                else:
                    return
            elif fmt == "替换特定字符":
                # 替换特定字符
                replace_dialog = QDialog(self)
                replace_dialog.setWindowTitle("替换字符")
                replace_layout = QVBoxLayout(replace_dialog)
                
                from_label = QLabel("要替换的字符:")
                from_edit = QLineEdit()
                replace_layout.addWidget(from_label)
                replace_layout.addWidget(from_edit)
                
                to_label = QLabel("替换为:")
                to_edit = QLineEdit()
                replace_layout.addWidget(to_label)
                replace_layout.addWidget(to_edit)
                
                preset_label = QLabel("或选择预设:")
                preset_combo = QComboBox()
                preset_combo.addItems(["无", "全角标点转半角", "半角标点转全角"])
                replace_layout.addWidget(preset_label)
                replace_layout.addWidget(preset_combo)
                
                btn_box4 = QHBoxLayout()
                ok_btn4 = QPushButton("确定")
                cancel_btn4 = QPushButton("取消")
                btn_box4.addStretch()
                btn_box4.addWidget(ok_btn4)
                btn_box4.addWidget(cancel_btn4)
                replace_layout.addLayout(btn_box4)
                
                user_confirmed4 = [False]
                from_text = [""]
                to_text = [""]
                preset_type = [0]
                
                def on_ok4():
                    from_text[0] = from_edit.text()
                    to_text[0] = to_edit.text()
                    preset_type[0] = preset_combo.currentIndex()
                    user_confirmed4[0] = True
                    replace_dialog.accept()
                
                ok_btn4.clicked.connect(on_ok4)
                cancel_btn4.clicked.connect(replace_dialog.reject)
                
                if replace_dialog.exec() and user_confirmed4[0]:
                    if preset_type[0] == 1:
                        # 全角转半角
                        formatted = self._full_to_half(text)
                    elif preset_type[0] == 2:
                        # 半角转全角
                        formatted = self._half_to_full(text)
                    else:
                        # 自定义替换
                        f = from_text[0]
                        t = to_text[0]
                        formatted = text.replace(f, t)
                else:
                    return
            elif fmt == "修复URL格式":
                # 自动检测并修复URL
                import re
                # 匹配URL的正则表达式
                url_pattern = r'(https?://[^\s]+|www\.[^\s]+)'
                
                def replace_url(match):
                    url = match.group(0)
                    # 确保URL有协议
                    if url.startswith('www.'):
                        url = 'https://' + url
                    # 返回可点击的Markdown链接格式
                    return f'[{url}]({url})'
                
                formatted = re.sub(url_pattern, replace_url, text)
            else:
                formatted = '\n'.join([l.strip() for l in text.split('\n')])
            
            self.output_text.setText(formatted)
            self.show_toast("✅ 格式化完成")
        except json.JSONDecodeError as e:
            self.show_toast(f"JSON解析失败: {str(e)}")
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            self.show_toast(f"格式化失败: {str(e)}")

    def _full_to_half(self, text):
        result = []
        for char in text:
            code = ord(char)
            if 65281 <= code <= 65374:
                result.append(chr(code - 65248))
            elif char == '\u3000':
                result.append(' ')
            else:
                result.append(char)
        return ''.join(result)

    def _half_to_full(self, text):
        result = []
        for char in text:
            code = ord(char)
            if 32 <= code <= 126:
                result.append(chr(code + 65248))
            elif char == ' ':
                result.append('\u3000')
            else:
                result.append(char)
        return ''.join(result)

    def _copy_formatted(self):
        text = self.output_text.toPlainText()
        if text:
            QApplication.clipboard().setText(text)
            self.show_toast("✅ 已复制到剪贴板")

    # ================== 8. 格式互转 ==================
    def open_format_convert(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("格式互转")
        dialog.setMinimumSize(520, 420)

        theme = {}
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            theme = self.main_window.theme_manager.get_theme()

        bg_color = theme.get("bg", "#ffffff")
        card_bg = theme.get("card", "#ffffff")
        text_color = theme.get("text", "#303133")
        border_color = theme.get("border", "#dcdfe6")
        primary_color = theme.get("primary", "#409eff")

        dialog.setStyleSheet(f"""
            QDialog {{
                background-color: {bg_color};
            }}
        """)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(20)

        title = QLabel("文档格式互转")
        title.setStyleSheet(f"font-size: 18px; font-weight: bold; color: {text_color};")
        layout.addWidget(title)

        info_label = QLabel("支持 Excel、Word、TXT、CSV 等多种格式相互转换\n点击按钮开始转换，大文件将显示处理进度")
        info_label.setStyleSheet(f"color: {text_color}; opacity: 0.7; font-size: 13px;")
        layout.addWidget(info_label)

        button_layout = QGridLayout()
        button_layout.setSpacing(12)
        button_layout.setRowStretch(0, 1)
        button_layout.setRowStretch(1, 1)
        button_layout.setColumnStretch(0, 1)
        button_layout.setColumnStretch(1, 1)

        class HoverButton(QPushButton):
            def __init__(self, text, normal_style, hover_style, parent=None):
                super().__init__(text, parent)
                self.normal_style = normal_style
                self.hover_style = hover_style
                self.setStyleSheet(normal_style)
                self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            def enterEvent(self, event):
                self.setStyleSheet(self.hover_style)
                super().enterEvent(event)
            def leaveEvent(self, event):
                self.setStyleSheet(self.normal_style)
                super().leaveEvent(event)

        normal_style = f"""
            padding: 20px;
            font-size: 15px;
            font-weight: 500;
            background: {card_bg};
            color: {text_color};
            border: 1px solid {border_color};
            border-radius: 12px;
            min-height: 80px;
        """
        hover_style = f"""
            padding: 20px;
            font-size: 15px;
            font-weight: 500;
            background: {primary_color};
            color: white;
            border: 1px solid {primary_color};
            border-radius: 12px;
            min-height: 80px;
        """

        conversions = [
            ("Excel转Word", self.start_excel_to_word_convert, "fa5s.file-excel"),
            ("Word转Excel", self.start_word_to_excel_convert, "fa5s.file-word"),
            ("Word转Txt", self.start_word_to_txt_convert, "fa5s.file-alt"),
            ("Excel转CSV", self.start_excel_to_csv_convert, "fa5s.file-csv"),
            ("Txt转Word", self.start_txt_to_word_convert, "fa5s.file-word"),
            ("CSV转Excel", self.start_csv_to_excel_convert, "fa5s.file-excel"),
        ]

        for i, (name, func, icon_name) in enumerate(conversions):
            btn = HoverButton(f"  {name}", normal_style, hover_style)
            btn.setIcon(qta.icon(icon_name, color=primary_color))
            btn.setIconSize(QSize(24, 24))
            btn.clicked.connect(func)
            button_layout.addWidget(btn, i // 2, i % 2)

        layout.addLayout(button_layout)
        layout.addStretch()
        dialog.exec()

    def convert_with_progress(self, task_func, task_args, title):
        theme = {}
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            theme = self.main_window.theme_manager.get_theme()
        progress_dialog = ConvertProgressDialog(title, self, theme)
        progress_dialog.start_convert(task_func, *task_args)
        progress_dialog.exec()

    def start_excel_to_word_convert(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择Excel", "", "Excel文件 (*.xlsx *.xls)")
        if not path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存Word", "", "Word文档 (*.docx)")
        if not save_path:
            return
        self.convert_with_progress(self._convert_excel_to_word, (path, save_path), "Excel转Word")

    def start_word_to_excel_convert(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择Word", "", "Word文档 (*.docx)")
        if not path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存Excel", "", "Excel文件 (*.xlsx)")
        if not save_path:
            return
        self.convert_with_progress(self._convert_word_to_excel, (path, save_path), "Word转Excel")

    def start_word_to_txt_convert(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择Word", "", "Word文档 (*.docx)")
        if not path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存Txt", "", "文本文件 (*.txt)")
        if not save_path:
            return
        self.convert_with_progress(self._convert_word_to_txt, (path, save_path), "Word转Txt")

    def start_excel_to_csv_convert(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择Excel", "", "Excel文件 (*.xlsx *.xls)")
        if not path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存CSV", "", "CSV文件 (*.csv)")
        if not save_path:
            return
        self.convert_with_progress(self._convert_excel_to_csv, (path, save_path), "Excel转CSV")
    
    def start_txt_to_word_convert(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择文本文件", "", "文本文件 (*.txt)")
        if not path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存Word", "", "Word文档 (*.docx)")
        if not save_path:
            return
        self.convert_with_progress(self._convert_txt_to_word, (path, save_path), "Txt转Word")
    
    def start_csv_to_excel_convert(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择CSV文件", "", "CSV文件 (*.csv)")
        if not path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "保存Excel", "", "Excel文件 (*.xlsx)")
        if not save_path:
            return
        self.convert_with_progress(self._convert_csv_to_excel, (path, save_path), "CSV转Excel")

    def _convert_excel_to_word(self, progress, log, thread, path, save_path):
        import openpyxl
        from docx import Document
        log.emit(f"正在打开文件: {os.path.basename(path)}")
        progress.emit(10)
        
        if thread.is_cancelled:
            return "已取消"
            
        wb = openpyxl.load_workbook(path)
        ws = wb.active
        if ws.max_row == 0 or ws.max_column == 0:
            return "Excel文件为空或没有数据"
        
        if thread.is_cancelled:
            return "已取消"
            
        log.emit(f"正在读取数据: {ws.max_row} 行 x {ws.max_column} 列")
        progress.emit(30)
        
        if thread.is_cancelled:
            return "已取消"
            
        doc = Document()
        table = doc.add_table(rows=1, cols=ws.max_column)
        table.style = 'Table Grid'
        log.emit("正在转换表头...")
        progress.emit(40)
        
        if thread.is_cancelled:
            return "已取消"
            
        for col in range(ws.max_column):
            table.rows[0].cells[col].text = str(ws.cell(1, col+1).value or "")
            
        total_rows = ws.max_row - 1
        for row_idx, row in enumerate(range(2, ws.max_row+1), 1):
            if thread.is_cancelled:
                return "已取消"
                
            cells = table.add_row().cells
            for col in range(ws.max_column):
                cells[col].text = str(ws.cell(row, col+1).value or "")
                
            if row_idx % 100 == 0:
                pct = 40 + int(50 * row_idx / max(total_rows, 1))
                progress.emit(pct)
                log.emit(f"已处理 {row_idx}/{total_rows} 行...")
                
        if thread.is_cancelled:
            return "已取消"
            
        progress.emit(90)
        log.emit("正在保存文件...")
        doc.save(save_path)
        progress.emit(100)
        log.emit(f"转换成功！\n保存位置: {save_path}")
        return "转换成功"

    def _convert_word_to_excel(self, progress, log, thread, path, save_path):
        from docx import Document
        import openpyxl
        log.emit(f"正在打开文件: {os.path.basename(path)}")
        progress.emit(10)
        
        if thread.is_cancelled:
            return "已取消"
            
        doc = Document(path)
        wb = openpyxl.Workbook()
        ws = wb.active
        row = 1
        has_content = False
        log.emit("正在读取文档内容...")
        progress.emit(30)
        
        if thread.is_cancelled:
            return "已取消"
            
        if doc.paragraphs:
            for p in doc.paragraphs:
                if thread.is_cancelled:
                    return "已取消"
                    
                if p.text.strip():
                    ws.cell(row, 1, value=p.text.strip())
                    row += 1
                    has_content = True
                    
        total_tables = len(doc.tables)
        for table_idx, table in enumerate(doc.tables, 1):
            if thread.is_cancelled:
                return "已取消"
                
            for r in table.rows:
                col = 1
                for cell in r.cells:
                    ws.cell(row, col, value=cell.text)
                    col += 1
                row += 1
                
            if table_idx % 5 == 0:
                pct = 30 + int(50 * table_idx / max(total_tables, 1))
                progress.emit(pct)
                log.emit(f"已处理 {table_idx}/{total_tables} 个表格...")
                
        if thread.is_cancelled:
            return "已取消"
            
        if not has_content and total_tables == 0:
            return "Word文档没有可转换的内容（表格或段落）"
            
        progress.emit(90)
        log.emit("正在保存文件...")
        wb.save(save_path)
        progress.emit(100)
        log.emit(f"转换成功！\n保存位置: {save_path}")
        return "转换成功"

    def _convert_word_to_txt(self, progress, log, thread, path, save_path):
        from docx import Document
        log.emit(f"正在打开文件: {os.path.basename(path)}")
        progress.emit(10)
        
        if thread.is_cancelled:
            return "已取消"
            
        doc = Document(path)
        log.emit("正在提取文字...")
        progress.emit(30)
        
        if thread.is_cancelled:
            return "已取消"
            
        text_parts = []
        total_paras = len(doc.paragraphs)
        for i, p in enumerate(doc.paragraphs):
            if thread.is_cancelled:
                return "已取消"
                
            if p.text.strip():
                text_parts.append(p.text)
                
            if i % 50 == 0 and total_paras > 0:
                pct = 30 + int(30 * i / total_paras)
                progress.emit(pct)
                
        total_tables = len(doc.tables)
        for table_idx, table in enumerate(doc.tables):
            if thread.is_cancelled:
                return "已取消"
                
            text_parts.append("\n--- 表格 ---\n")
            for r in table.rows:
                row_text = " | ".join([cell.text for cell in r.cells])
                text_parts.append(row_text)
            text_parts.append("\n")
            
            if table_idx % 5 == 0 and total_tables > 0:
                pct = 60 + int(30 * table_idx / total_tables)
                progress.emit(pct)
                
        if thread.is_cancelled:
            return "已取消"
            
        progress.emit(90)
        log.emit("正在保存文件...")
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_parts))
        progress.emit(100)
        log.emit(f"转换成功！\n保存位置: {save_path}")
        return "转换成功"

    def _convert_excel_to_csv(self, progress, log, thread, path, save_path):
        import openpyxl
        import csv
        log.emit(f"正在打开文件: {os.path.basename(path)}")
        progress.emit(10)
        
        if thread.is_cancelled:
            return "已取消"
            
        wb = openpyxl.load_workbook(path)
        ws = wb.active
        log.emit(f"正在读取数据: {ws.max_row} 行 x {ws.max_column} 列")
        progress.emit(30)
        
        if thread.is_cancelled:
            return "已取消"
            
        progress.emit(50)
        log.emit("正在转换...")
        with open(save_path, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            total_rows = ws.max_row
            for i, row in enumerate(ws.iter_rows(values_only=True), 1):
                if thread.is_cancelled:
                    return "已取消"
                    
                writer.writerow(row)
                
                if i % 100 == 0:
                    pct = 50 + int(40 * i / max(total_rows, 1))
                    progress.emit(pct)
                    log.emit(f"已处理 {i}/{total_rows} 行...")
                    
        if thread.is_cancelled:
            return "已取消"
            
        progress.emit(95)
        log.emit("正在保存文件...")
        progress.emit(100)
        log.emit(f"转换成功！\n保存位置: {save_path}")
        return "转换成功"
    
    def _convert_txt_to_word(self, progress, log, thread, path, save_path):
        from docx import Document
        log.emit(f"正在读取文件: {os.path.basename(path)}")
        progress.emit(10)
        
        if thread.is_cancelled:
            return "已取消"
            
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            
        progress.emit(40)
        log.emit("正在创建Word文档...")
        
        if thread.is_cancelled:
            return "已取消"
            
        doc = Document()
        paragraphs = content.split('\n')
        total = len(paragraphs)
        for i, para in enumerate(paragraphs):
            if thread.is_cancelled:
                return "已取消"
                
            if para.strip():
                doc.add_paragraph(para)
                
            if i % 200 == 0:
                pct = 40 + int(50 * i / max(total, 1))
                progress.emit(pct)
                
        if thread.is_cancelled:
            return "已取消"
            
        progress.emit(90)
        log.emit("正在保存文件...")
        doc.save(save_path)
        progress.emit(100)
        log.emit(f"转换成功！\n保存位置: {save_path}")
        return "转换成功"
    
    def _convert_csv_to_excel(self, progress, log, thread, path, save_path):
        import csv
        import openpyxl
        log.emit(f"正在读取文件: {os.path.basename(path)}")
        progress.emit(10)
        
        if thread.is_cancelled:
            return "已取消"
            
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        with open(path, 'r', encoding='utf-8-sig', errors='ignore') as f:
            reader = csv.reader(f)
            row_idx = 1
            total_rows = 0
            for row in reader:
                total_rows += 1
            f.seek(0)
            
            for i, row in enumerate(reader):
                if thread.is_cancelled:
                    return "已取消"
                    
                for col_idx, value in enumerate(row, 1):
                    ws.cell(row_idx, col_idx, value=value)
                row_idx += 1
                
                if i % 500 == 0:
                    pct = 10 + int(80 * i / max(total_rows, 1))
                    progress.emit(pct)
                    log.emit(f"已处理 {i}/{total_rows} 行...")
                    
        if thread.is_cancelled:
            return "已取消"
            
        progress.emit(90)
        log.emit("正在保存文件...")
        wb.save(save_path)
        progress.emit(100)
        log.emit(f"转换成功！\n保存位置: {save_path}")
        return "转换成功"
    
    def convert_excel_to_word(self):
        self.start_excel_to_word_convert()

    def convert_word_to_excel(self):
        self.start_word_to_excel_convert()

    def convert_word_to_txt(self):
        self.start_word_to_txt_convert()

    def convert_excel_to_csv(self):
        self.start_excel_to_csv_convert()

    # ================== 9. 音频转换 ==================
    def open_audio_convert(self):
        if not self._check_pro_or_warn():
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("音频格式转换")
        dialog.setMinimumSize(600, 450)
        layout = QVBoxLayout(dialog)

        file_list_label = QLabel("已选音视频文件：")
        layout.addWidget(file_list_label)
        self.audio_file_list = QListWidget()
        layout.addWidget(self.audio_file_list)

        btn_row = QHBoxLayout()
        select_btn = QPushButton(" 选择音视频文件")
        select_btn.setIcon(qta.icon('fa5s.folder-open', color='white'))
        select_btn.setIconSize(ICON_SIZE_SMALL)
        select_btn.clicked.connect(self._select_audio_files)
        btn_row.addWidget(select_btn)

        clear_btn = QPushButton(" 清空列表")
        clear_btn.setIcon(qta.icon('fa5s.trash-alt', color='white'))
        clear_btn.clicked.connect(lambda: self.audio_file_list.clear())
        btn_row.addWidget(clear_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        format_label = QLabel("输出格式：")
        layout.addWidget(format_label)
        self.audio_format_combo = QComboBox()
        self.audio_format_combo.addItems(["MP3", "WAV", "FLAC", "M4A", "AAC", "OGG", "OPUS"])
        layout.addWidget(self.audio_format_combo)

        output_label = QLabel("输出目录：")
        layout.addWidget(output_label)
        output_row = QHBoxLayout()
        self.audio_output_path = QLineEdit()
        self.audio_output_path.setPlaceholderText("选择一个文件夹保存转换后的音频...")
        output_row.addWidget(self.audio_output_path)
        select_output_btn = QPushButton(" 浏览")
        select_output_btn.clicked.connect(self._select_audio_output)
        output_row.addWidget(select_output_btn)
        layout.addLayout(output_row)

        convert_btn = QPushButton(" 开始转换")
        convert_btn.setIcon(qta.icon('fa5s.music', color='white'))
        convert_btn.clicked.connect(self._start_audio_convert_with_progress)
        layout.addWidget(convert_btn)

        hint_label = QLabel("支持格式: MP3, WAV, FLAC, M4A, AAC, OGG, OPUS | 支持视频提取音频")
        hint_label.setStyleSheet("color: #6b7280; font-size: 11px; padding: 8px; background-color: #f3f4f6; border-radius: 4px;")
        layout.addWidget(hint_label)

        dialog.exec()

    def _select_audio_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择音视频文件", "",
            "音视频文件 (*.mp3 *.wav *.flac *.m4a *.aac *.ogg *.wma *.mp4 *.mkv *.avi *.mov *.flv *.MP3 *.WAV *.FLAC *.M4A *.AAC *.OGG *.WMA *.MP4 *.MKV *.AVI *.MOV *.FLV)")
        for f in files:
            self.audio_file_list.addItem(f)

    def _select_audio_output(self):
        folder = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if folder:
            self.audio_output_path.setText(folder)

    def _start_audio_convert_with_progress(self):
        if self.audio_file_list.count() == 0:
            self.show_toast("请先选择要转换的音频文件")
            return
        output_folder = self.audio_output_path.text()
        if not output_folder:
            self.show_toast("请选择输出目录")
            return
        output_format = self.audio_format_combo.currentText().lower()
        
        ffmpeg_path = 'ffmpeg'
        import shutil
        if not shutil.which('ffmpeg'):
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            ffmpeg_exe = os.path.join(base_dir, 'ffmpeg-8.0.1', 'bin', 'ffmpeg.exe')
            if os.path.exists(ffmpeg_exe):
                ffmpeg_path = ffmpeg_exe
            else:
                ffmpeg_exe = os.path.join(base_dir, 'ffmpeg', 'ffmpeg.exe')
                if os.path.exists(ffmpeg_exe):
                    ffmpeg_path = ffmpeg_exe
                else:
                    self.show_warning(
                        "缺少依赖",
                        "未找到FFmpeg。\n\n请：\n1. 安装FFmpeg并添加到PATH\n2. 或在设置中指定FFmpeg路径\n\n下载地址：https://ffmpeg.org/download.html"
                    )
                    return
        
        file_list = [self.audio_file_list.item(i).text() for i in range(self.audio_file_list.count())]
        
        theme = {}
        if self.main_window and hasattr(self.main_window, 'theme_manager'):
            theme = self.main_window.theme_manager.get_theme()
        
        progress_dialog = ConvertProgressDialog("音频格式转换", self, theme)
        
        self._audio_thread = AudioThread(file_list, output_folder, output_format, ffmpeg_path)
        self._audio_thread.progress.connect(progress_dialog.update_progress)
        self._audio_thread.log.connect(progress_dialog.add_log)
        self._audio_thread.finished.connect(lambda success, msg: self._on_audio_convert_finished(success, msg, progress_dialog))
        progress_dialog.cancel_btn.clicked.connect(lambda: self._cancel_audio_convert(progress_dialog))
        self._audio_thread.start()
        progress_dialog.exec()

    def _on_audio_convert_finished(self, success, message, progress_dialog):
        progress_dialog.progress_bar.setValue(100 if success else 0)
        progress_dialog.status_label.setText("转换完成" if success else "转换失败")
        progress_dialog.status_label.setStyleSheet(
            "color: #67c23a; font-size: 13px;" if success else "color: #f56c6c; font-size: 13px;"
        )
        progress_dialog.add_log(message)
        progress_dialog.cancel_btn.setText("关闭")
        self.show_toast(message)
        self._audio_thread = None

    def _cancel_audio_convert(self, progress_dialog):
        if hasattr(self, '_audio_thread') and self._audio_thread:
            self._audio_thread.cancel()
            self._audio_thread.wait(1000)
            self._audio_thread = None
        progress_dialog.is_running = False

    # ================== 10. 截图功能 ==================
    def open_screenshot(self, mode='full'):
        if mode == 'region':
            self._start_region_screenshot()
        elif mode == 'window':
            self._start_window_screenshot()
        else:
            self._start_fullscreen_screenshot()

    def _start_fullscreen_screenshot(self):
        from PySide6.QtWidgets import QApplication
        screens = QApplication.screens()
        if screens:
            total_geometry = QRect()
            for s in screens:
                total_geometry = total_geometry.united(s.geometry())
            combined_pixmap = QPixmap(total_geometry.size())
            combined_pixmap.fill(Qt.black)
            painter = QPainter(combined_pixmap)
            for screen in screens:
                screen_pixmap = screen.grabWindow(0)
                painter.drawPixmap(screen.geometry().topLeft(), screen_pixmap)
            painter.end()
            self._show_screenshot_result(combined_pixmap)
        else:
            self.show_toast("无法获取屏幕信息")

    def _start_window_screenshot(self):
        if self.main_window and self.main_window.isVisible():
            self.main_window.hide()
        QTimer.singleShot(300, lambda: self._show_window_region_selector())

    def _show_window_region_selector(self):
        selector = RegionSelector(None)  # 设置父窗口为None，让区域选择器覆盖整个屏幕
        selector.setWindowTitle("选择窗口区域")
        if selector.exec() == QDialog.Accepted:
            screenshot = selector.get_screenshot()
            self._show_screenshot_result(screenshot)

    def _start_region_screenshot(self):
        if self.main_window and self.main_window.isVisible():
            self.main_window.hide()
        QTimer.singleShot(300, lambda: self._show_region_selector())

    def _show_region_selector(self):
        selector = RegionSelector(None)  # 设置父窗口为None，让区域选择器覆盖整个屏幕
        if selector.exec() == QDialog.Accepted:
            screenshot = selector.get_screenshot()
            self._show_screenshot_result(screenshot)

    def _show_screenshot_result(self, screenshot):
        if screenshot.isNull():
            self.show_toast("截图失败")
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("截图结果")
        dialog.setMinimumSize(600, 500)
        layout = QVBoxLayout(dialog)
        label = QLabel()
        label.setAlignment(Qt.AlignCenter)
        label.setStyleSheet("border: 1px solid #ddd; background: #f5f5f5;")
        scaled_pixmap = screenshot.scaled(550, 400, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        label.setPixmap(scaled_pixmap)
        layout.addWidget(label)
        btn_layout = QHBoxLayout()
        save_btn = QPushButton(" 保存到文件")
        save_btn.setIcon(qta.icon('fa5s.save', color='white'))
        save_btn.clicked.connect(lambda: self._save_screenshot(screenshot, dialog))
        btn_layout.addWidget(save_btn)
        copy_btn = QPushButton(" 复制到剪贴板")
        copy_btn.setIcon(qta.icon('fa5s.copy', color='white'))
        copy_btn.clicked.connect(lambda: self._copy_screenshot(screenshot, dialog))
        btn_layout.addWidget(copy_btn)
        close_btn = QPushButton(" 关闭")
        close_btn.setIcon(qta.icon('fa5s.times', color='white'))
        close_btn.clicked.connect(dialog.accept)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        dialog.exec()

    def _save_screenshot(self, screenshot, parent_dialog):
        save_path, _ = QFileDialog.getSaveFileName(parent_dialog, "保存截图", "", "PNG图片 (*.png);;JPG图片 (*.jpg);;所有文件 (*.*)")
        if save_path:
            if screenshot.save(save_path):
                self.show_toast("✅ 保存成功")
            else:
                self.show_toast("❌ 保存失败")

    def _copy_screenshot(self, screenshot, parent_dialog):
        clipboard = QApplication.clipboard()
        clipboard.setPixmap(screenshot)
        self.show_toast("✅ 已复制到剪贴板")
        parent_dialog.accept()


class RegionSelector(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)

        from PySide6.QtWidgets import QApplication
        screens = QApplication.screens()
        if screens:
            screen = screens[0]
            screen_geometry = screen.geometry()
            self.setGeometry(screen_geometry)
            self.full_screen = screen.grabWindow(0)
        else:
            self.setGeometry(0, 0, 1920, 1080)
            self.full_screen = QPixmap(1920, 1080)

        self.screenshot = None
        self.start_pos = None
        self.end_pos = None
        self.selection = None
        self.is_selecting = False
        self.setCursor(Qt.CrossCursor)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.drawPixmap(0, 0, self.full_screen)
        painter.fillRect(self.rect(), QColor(0, 0, 0, 128))
        if self.selection:
            painter.setCompositionMode(QPainter.CompositionMode_Clear)
            painter.fillRect(self.selection, QColor(0, 0, 0, 0))
            painter.setCompositionMode(QPainter.CompositionMode_SourceOver)
            pen = QPen(QColor(255, 0, 0), 2)
            painter.setPen(pen)
            painter.drawRect(self.selection)
            w = self.selection.width()
            h = self.selection.height()
            info_text = f"{w} x {h}"
            painter.setPen(Qt.white)
            painter.drawText(self.selection.topLeft().x() + 5, self.selection.topLeft().y() - 5, info_text)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.is_selecting = True
            self.start_pos = event.pos()
            self.end_pos = event.pos()
            self.selection = QRect(self.start_pos, self.end_pos)
            self.update()

    def mouseMoveEvent(self, event):
        if self.is_selecting:
            self.end_pos = event.pos()
            self.selection = QRect(self.start_pos, self.end_pos).normalized()
            self.update()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton and self.is_selecting:
            self.is_selecting = False
            if self.selection.width() > 5 and self.selection.height() > 5:
                self.screenshot = self.full_screen.copy(self.selection)
                self.accept()
            else:
                self.reject()

    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Escape:
            self.reject()

    def get_screenshot(self):
        return self.screenshot